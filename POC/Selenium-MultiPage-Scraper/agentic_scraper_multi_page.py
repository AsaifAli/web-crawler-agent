import os
import base64
import sqlite3
import argparse
import time
import asyncio
from datetime import datetime
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager  # Added this import
from docx import Document
from markdown import markdown
import subprocess
from selenium.webdriver.common.action_chains import ActionChains
from concurrent.futures import ThreadPoolExecutor
import logging

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("scraper.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

DB_PATH = "crawl_metadata.db"

def setup_db():
    """Initialize SQLite database to store crawl metadata"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS crawl_data (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                url TEXT,
                timestamp TEXT,
                directory TEXT,
                summary TEXT
            )
        """)
        conn.commit()
        conn.close()
        logger.info(f"Database initialized at {DB_PATH}")
    except sqlite3.Error as e:
        logger.error(f"Database error: {e}")
        raise

def save_metadata(url, timestamp, directory, summary):
    """Save crawl metadata to SQLite database"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO crawl_data (url, timestamp, directory, summary) VALUES (?, ?, ?, ?)",
                (url, timestamp, directory, summary))
        conn.commit()
        conn.close()
        logger.info(f"Metadata saved for {url}")
    except sqlite3.Error as e:
        logger.error(f"Error saving metadata: {e}")

def extract_links(html, base_url):
    """Extract all links from HTML that belong to the same domain"""
    soup = BeautifulSoup(html, "html.parser")
    links = set()
    for a_tag in soup.find_all("a", href=True):
        href = a_tag.get("href")
        full_url = urljoin(base_url, href)
        if urlparse(full_url).netloc == urlparse(base_url).netloc:
            links.add(full_url)
    logger.info(f"Extracted {len(links)} same-domain links from {base_url}")
    return list(links)

def setup_selenium_driver():
    """Set up Chrome WebDriver with webdriver-manager for automatic chromedriver management"""
    try:
        options = Options()
        options.add_argument("--headless")  # Run Chrome in headless mode
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--window-size=1920,1080")
        
        # Use ChromeDriverManager to automatically download and manage chromedriver
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        return driver
    except Exception as e:
        logger.error(f"Error setting up Chrome driver: {e}")
        raise

def crawl_and_process(url, output_root, visited):
    """Crawl a URL and process its content"""
    if url in visited:
        logger.info(f"URL already visited: {url}")
        return None, None

    logger.info(f"🌐 Crawling: {url}")
    visited.add(url)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_url = url.replace("https://", "").replace("http://", "").replace("/", "_").replace(":", "_")
    site_dir = os.path.join(output_root, f"{safe_url}_{timestamp}")
    
    try:
        os.makedirs(site_dir, exist_ok=True)
    except Exception as e:
        logger.error(f"Error creating directory {site_dir}: {e}")
        return None, None

    try:
        # Setup Selenium WebDriver
        driver = setup_selenium_driver()
        driver.get(url)
        # Wait for page to load
        time.sleep(2)
        page_html = driver.page_source

        # Generate markdown and HTML outputs
        md_path = os.path.join(site_dir, "index.md")
        html_path = os.path.join(site_dir, "index.html")
        
        # Convert HTML to markdown first
        html_text = page_html
        md_content = html_to_markdown(html_text)
        
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(page_html)

        # Capture screenshot
        driver.save_screenshot(os.path.join(site_dir, "screenshot.png"))

        # Extract and save information about forms
        soup = BeautifulSoup(page_html, "html.parser")
        forms = soup.find_all("form")
        forms_summary = []
        for form in forms:
            inputs = form.find_all("input")
            forms_summary.append({
                "action": form.get("action"),
                "method": form.get("method"),
                "inputs": [{"name": i.get("name"), "type": i.get("type")} for i in inputs]
            })

        with open(os.path.join(site_dir, "forms.txt"), "w", encoding="utf-8") as f:
            for form in forms_summary:
                f.write(str(form) + "\n")

        driver.quit()
        logger.info(f"Successfully processed {url}")
        return md_content, site_dir
        
    except Exception as e:
        logger.error(f"Error crawling {url}: {e}")
        if 'driver' in locals():
            driver.quit()
        return None, None

def html_to_markdown(html_text):
    """Convert HTML to markdown using BeautifulSoup for better control"""
    try:
        soup = BeautifulSoup(html_text, "html.parser")
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        # Get text
        text = soup.get_text(separator="\n", strip=True)
        return text
    except Exception as e:
        logger.error(f"Error converting HTML to markdown: {e}")
        return markdown(html_text)  # Fallback to markdown library

def export_to_docx(markdown_text, site_dir):
    """Export markdown text to DOCX format"""
    try:
        doc = Document()
        doc.add_heading("Web Summary", 0)
        for line in markdown_text.split("\n"):
            if line.strip():  # Skip empty lines
                doc.add_paragraph(line)
        doc_path = os.path.join(site_dir, "index.docx")
        doc.save(doc_path)
        logger.info(f"DOCX exported to {doc_path}")
    except Exception as e:
        logger.error(f"Error exporting to DOCX: {e}")

def summarize_with_ollama(markdown_text, site_dir):
    """Generate a summary using Ollama LLM"""
    try:
        if len(markdown_text) > 8000:
            # Truncate text if too long
            markdown_text = markdown_text[:8000] + "..."
        
        prompt = f"Summarize the following website content in 3-5 sentences:\n\n{markdown_text}"
        result = subprocess.run(
            ["ollama", "run", "llama3.2:1b", prompt],
            capture_output=True, text=True, timeout=60  # Added timeout
        )
        summary = result.stdout.strip()
        
        if not summary:
            summary = "Summary generation failed."
            logger.warning("Ollama summary generation returned empty result")
        
        summary_path = os.path.join(site_dir, "summary.txt")
        with open(summary_path, "w", encoding="utf-8") as f:
            f.write(summary)
        logger.info(f"Summary created at {summary_path}")
        return summary
    except subprocess.TimeoutExpired:
        logger.error("Ollama summary generation timed out")
        return "Summary generation timed out."
    except Exception as e:
        logger.error(f"Error generating summary: {e}")
        return f"Error generating summary: {str(e)}"

async def main_async(urls, interval, output_root, max_depth=1, max_pages_per_domain=10):
    """Asynchronous main function to manage crawling"""
    setup_db()
    
    while True:
        for root_url in urls:
            visited = set()
            to_crawl = [root_url]
            pages_crawled = 0
            
            for depth in range(max_depth):
                logger.info(f"Crawling depth {depth+1}/{max_depth} for {root_url}")
                new_links = []
                
                with ThreadPoolExecutor(max_workers=5) as executor:
                    futures = []
                    for url in to_crawl:
                        if pages_crawled >= max_pages_per_domain:
                            break
                        futures.append(executor.submit(crawl_and_process, url, output_root, visited))
                        pages_crawled += 1
                    
                    for future in futures:
                        markdown_text, site_dir = future.result()
                        if markdown_text and site_dir:
                            export_to_docx(markdown_text, site_dir)
                            summary = summarize_with_ollama(markdown_text, site_dir)
                            save_metadata(url, datetime.now().strftime("%Y%m%d_%H%M%S"), site_dir, summary)
                            
                            # Extract links for next depth
                            links = extract_links(markdown_text, url)
                            with open(os.path.join(site_dir, "links.txt"), "w", encoding="utf-8") as f:
                                for link in links:
                                    f.write(link + "\n")
                                    if link not in visited and pages_crawled < max_pages_per_domain:
                                        new_links.append(link)
                
                if not new_links or pages_crawled >= max_pages_per_domain:
                    break
                    
                to_crawl = new_links

        if interval <= 0:
            break

        # Sleep between crawls
        logger.info(f"Sleeping for {interval} hours before next crawl")
        await asyncio.sleep(interval * 3600)

def main(urls, interval, output_root, max_depth=1, max_pages=10):
    """Main function that sets up and runs the async loop"""
    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
    
    try:
        loop.run_until_complete(main_async(urls, interval, output_root, max_depth, max_pages))
    except KeyboardInterrupt:
        logger.info("Crawling stopped by user")
    except Exception as e:
        logger.error(f"Crawling error: {e}")
    finally:
        loop.close()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Multi-page Web Scraper with Selenium and DB support")
    parser.add_argument("--url", nargs="+", required=True, help="List of URLs to crawl")
    parser.add_argument("--interval", type=int, default=0, help="Crawl interval in hours (0 to disable loop)")
    parser.add_argument("--output", type=str, default="output/sites", help="Root output directory")
    parser.add_argument("--depth", type=int, default=1, help="Maximum crawl depth")
    parser.add_argument("--max-pages", type=int, default=10, help="Maximum pages to crawl per domain")

    args = parser.parse_args()
    logger.info(f"Starting crawler with URLs: {args.url}")
    main(args.url, args.interval, args.output, args.depth, args.max_pages)