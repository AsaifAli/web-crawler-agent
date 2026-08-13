import os
import json
import logging
from datetime import datetime
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
import markdown2
from PIL import Image
import io
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.schema import StrOutputParser

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ImageContent(BaseModel):
    """Pydantic model for structured image content."""
    filename: str
    timestamp: str
    extracted_text: str
    summary: str

class ImageToDocProcessor:
    """A processor to scan local images, extract text, and create functional documents."""

    def __init__(self, image_folder="screenshots", app_id=os.getenv("POC_APP_ID", ""), gemini_api_key=""):
        """Initialize the processor."""
        self.image_folder = image_folder
        self.app_id = app_id
        self.gemini_api_key = gemini_api_key
        self.image_details = {}
        self.report_dir = "reports"
        self.debug_dir = "debug"

        # Configure Gemini API
        if self.gemini_api_key:
            try:
                genai.configure(api_key=self.gemini_api_key)
            except Exception as e:
                logger.error(f"Error configuring Gemini API: {e}")
                self.gemini_api_key = ""

        # Initialize LangChain
        if self.gemini_api_key:
            try:
                self.llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    google_api_key=self.gemini_api_key,
                    max_output_tokens=1000
                )
            except Exception as e:
                logger.error(f"Error initializing LangChain LLM: {e}")
                self.llm = None
        else:
            self.llm = None

        # Create directories
        for directory in [self.image_folder, self.report_dir, self.debug_dir]:
            os.makedirs(directory, exist_ok=True)

    def extract_text_from_image(self, image_path: str) -> str:
        """Extract text from an image using Gemini."""
        if not self.gemini_api_key:
            logger.warning(f"Vision-based extraction skipped for {image_path}: Missing Gemini API key")
            return ""
        try:
            with Image.open(image_path) as img:
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_data = img_byte_arr.getvalue()
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(
                [
                    "Extract all visible text from this image. Return only the text.",
                    {"mime_type": "image/png", "data": img_data}
                ]
            )
            text = response.text.strip()
            logger.info(f"Vision-based text extracted from {image_path}: {len(text)} characters")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from image {image_path}: {e}")
            return ""

    def summarize_content(self, content: str) -> str:
        """Summarize content using LangChain."""
        if not self.llm:
            logger.warning("LangChain LLM not initialized, skipping summarization")
            return content[:200] + "..." if content else ""
        try:
            prompt = PromptTemplate(
                input_variables=["content"],
                template="Summarize the following text in 2-3 sentences, focusing on the main purpose and key features:\n{content}"
            )
            chain = prompt | self.llm | StrOutputParser()
            summary = chain.invoke({"content": content[:5000]})
            logger.info("Generated summary with LangChain")
            return summary
        except Exception as e:
            logger.error(f"Error summarizing content with LangChain: {e}")
            return content[:200] + "..." if content else ""

    def scan_images(self):
        """Scan images in the folder and extract text."""
        logger.info(f"Scanning images in {self.image_folder}")
        image_files = [f for f in os.listdir(self.image_folder) if f.lower().endswith('.png')]
        if not image_files:
            logger.warning(f"No .png images found in {self.image_folder}")
            return

        for image_file in image_files:
            image_path = os.path.join(self.image_folder, image_file)
            try:
                # Extract timestamp from filename (e.g., <app_id>_main_initial_<timestamp>.png)
                timestamp = image_file.split('_')[-2] + '_' + image_file.split('_')[-1].replace('.png', '')
                extracted_text = self.extract_text_from_image(image_path)
                summary = self.summarize_content(extracted_text) if extracted_text else ""

                self.image_details[image_file] = {
                    "filename": image_file,
                    "timestamp": timestamp,
                    "extracted_text": extracted_text,
                    "summary": summary,
                    "image_path": image_path
                }
                logger.info(f"Processed image: {image_file}")
            except Exception as e:
                logger.error(f"Error processing image {image_file}: {e}")

    def generate_report(self):
        """Generate Markdown, HTML, and DOCX reports."""
        logger.info("Generating reports...")
        report_content = f"# Image Processing Report\n\n"
        report_content += f"**Generated on**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report_content += f"**Application ID**: {self.app_id}\n"
        report_content += f"**Image Folder**: {self.image_folder}\n\n"
        report_content += "## Processed Images\n\n"

        for image_file, details in self.image_details.items():
            report_content += f"### Image: {image_file}\n"
            report_content += f"- **Timestamp**: {details['timestamp']}\n"
            if details['extracted_text']:
                report_content += f"- **Extracted Text**: {details['extracted_text'][:200]}... ({len(details['extracted_text'])} characters)\n"
            if details['summary']:
                report_content += f"- **Summary**: {details['summary']}\n"
            report_content += f"- **Image Path**: ![{image_file}]({details['image_path']})\n\n"

        # Save Markdown report
        report_path = os.path.join(self.report_dir, f"{self.app_id}_image_report.md")
        try:
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            logger.info(f"Report saved: {report_path}")
        except Exception as e:
            logger.error(f"Error saving Markdown report: {e}")

        # Save HTML report
        html_content = markdown2.markdown(report_content)
        html_path = os.path.join(self.report_dir, f"{self.app_id}_image_report.html")
        try:
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(f"""
<!DOCTYPE html>
<html>
<head>
    <title>Image Processing Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
""")
            logger.info(f"HTML Report saved: {html_path}")
        except Exception as e:
            logger.error(f"Error saving HTML report: {e}")

        # Save DOCX report
        doc = Document()
        doc.add_heading("Image Processing Report", 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Application ID: {self.app_id}")
        doc.add_paragraph(f"Image Folder: {self.image_folder}")
        doc.add_heading("Processed Images", level=1)

        for image_file, details in self.image_details.items():
            doc.add_heading(f"Image: {image_file}", level=2)
            doc.add_paragraph(f"Timestamp: {details['timestamp']}")
            if details['extracted_text']:
                doc.add_paragraph(f"Extracted Text: {details['extracted_text'][:200]}... ({len(details['extracted_text'])} characters)")
            if details['summary']:
                doc.add_paragraph(f"Summary: {details['summary']}")
            if os.path.exists(details['image_path']):
                try:
                    doc.add_picture(details['image_path'], width=Inches(6))
                except Exception as e:
                    logger.error(f"Error adding image {image_file} to DOCX: {e}")

        docx_path = os.path.join(self.report_dir, f"{self.app_id}_image_report.docx")
        try:
            doc.save(docx_path)
            logger.info(f"DOCX report saved: {docx_path}")
        except Exception as e:
            logger.error(f"Error saving DOCX report: {e}")
            docx_path = ""

        return report_path, html_path, docx_path

    def process(self):
        """Execute the image processing and report generation."""
        logger.info("Running ImageToDocProcessor for local image scanning and document creation")
        try:
            self.scan_images()
            if not self.image_details:
                logger.warning("No images processed, skipping report generation")
                return {}
            report_path, html_path, docx_path = self.generate_report()
            result = {
                "image_details": self.image_details,
                "report_path": report_path,
                "html_report_path": html_path,
                "docx_report_path": docx_path
            }
            with open('image_processing_results.json', 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2)
            logger.info("Image processing completed successfully!")
            logger.info(f"Images processed: {len(self.image_details)}")
            logger.info(f"Report saved: {report_path}")
            logger.info(f"HTML Report saved: {html_path}")
            logger.info(f"DOCX Report saved: {docx_path}")
            return result
        except Exception as e:
            logger.error(f"Error during image processing: {e}")
            raise

if __name__ == "__main__":
    processor = ImageToDocProcessor(
        image_folder="<folder>",
        app_id="<id>",
        gemini_api_key=os.getenv("GEMINI_API_KEY", "")
    )
    try:
        result = processor.process()
    except Exception as e:
        logger.error(f"Error running ImageToDocProcessor: {e}")