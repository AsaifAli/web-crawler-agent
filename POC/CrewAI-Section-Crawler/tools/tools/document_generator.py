# Reconstructed during backup cleanup — not present anywhere in the backup. Written to
# match its instantiation `DocumentGenerator(self.app_id, current_application_url)` in
# crawlerservice.py and the document_creator_task's `output_pydantic=DocumentResponseModel`
# (which only has a `directory: str` field) — so this tool's job is to write a document
# somewhere and hand back the directory it landed in.
from __future__ import annotations

import os
from datetime import datetime
from typing import List, Optional, Type

from pydantic import BaseModel, Field
from crewai.tools import BaseTool


class DocumentGeneratorInput(BaseModel):
    title: str = Field(..., description="Document title")
    content: str = Field(..., description="Document body content (plain text / markdown-ish)")


class DocumentGenerator(BaseTool):
    """Write a crawl summary document (docx if python-docx is available, else .txt) to disk."""

    name: str = "generate_document"
    description: str = "Generate a document from crawled/analyzed content and save it to disk."
    args_schema: Type[BaseModel] = DocumentGeneratorInput

    app_id: Optional[str] = Field(None)
    application_url: Optional[str] = Field(None)
    output_dir: str = Field(default="")

    def __init__(self, app_id: Optional[str] = None, application_url: Optional[str] = None, **kwargs):
        super().__init__(**kwargs)
        self.app_id = app_id
        self.application_url = application_url
        base = os.path.join(os.path.dirname(os.path.dirname(__file__)), "runs", "documents")
        self.output_dir = os.path.join(base, str(app_id or "default"))

    def _run(self, title: str, content: str, run_manager: Optional[object] = None) -> str:
        os.makedirs(self.output_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        try:
            from docx import Document

            doc = Document()
            doc.add_heading(title, level=1)
            if self.application_url:
                doc.add_paragraph(f"Source: {self.application_url}")
            for paragraph in content.split("\n\n"):
                doc.add_paragraph(paragraph)
            path = os.path.join(self.output_dir, f"{timestamp}_{title[:40]}.docx")
            doc.save(path)
        except ImportError:
            path = os.path.join(self.output_dir, f"{timestamp}_{title[:40]}.txt")
            with open(path, "w", encoding="utf-8") as f:
                f.write(f"{title}\n{'=' * len(title)}\n\n")
                if self.application_url:
                    f.write(f"Source: {self.application_url}\n\n")
                f.write(content)

        return self.output_dir
