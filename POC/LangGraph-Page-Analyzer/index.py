import os
import uuid
import json
import asyncio
import logging
from typing import List, Dict, TypedDict
from urllib.parse import urljoin, urlparse, urldefrag
from docx import Document
from pydantic import BaseModel, Field
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError, Error as PlaywrightError
from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig, CacheMode, LLMConfig
from crawl4ai.extraction_strategy import LLMExtractionStrategy, JsonCssExtractionStrategy
from langchain_ollama import ChatOllama
from langgraph.graph import StateGraph, END
from tqdm.asyncio import tqdm_asyncio

# Configure logging
logging.basicConfig(level=logging.INFO, filename="crawler.log", format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Configuration
LLM_MODEL = os.getenv("OLLAMA_MODEL", "")
MAX_DEPTH = 9
MAX_URLS = 1000
MAX_TOKENS_CONTENT = 3000  # Approx ~3000 words
OUTPUT_DOC = "website_analysis_langgraph.docx"
WAIT_UNTIL = "networkidle"
TIMEOUT = 60000
CHUNK_TOKEN_THRESHOLD = 1000
OVERLAP_RATE = 0.1

# Pydantic models
class Link(BaseModel):
    url: str = Field(..., description="The URL of the link")
    text: str = Field(..., description="Anchor text of the link")

class Form(BaseModel):
    action: str = Field(..., description="Form action URL")
    method: str = Field(..., description="Form method (e.g., GET, POST)")
    inputs: List[Dict] = Field(..., description="List of input fields")
    is_login: bool = Field(False, description="Is this a login form")

class Section(BaseModel):
    tag: str = Field(..., description="HTML tag (e.g., section, nav)")
    class_name: str = Field("", description="CSS class name")
    id: str = Field("", description="Element ID")

class Widget(BaseModel):
    type: str = Field(..., description="Widget type (e.g., slider, modal, iframe)")
    selector: str = Field(..., description="CSS selector for the widget")

class PageElements(BaseModel):
    forms: List[Form] = Field(..., description="List of forms on the page")
    sections: List[Section] = Field(..., description="List of sections on the page")
    widgets: List[Widget] = Field(..., description="List of widgets on the page")

class PageAnalysis(BaseModel):
    url: str
    title: str
    summary: str
    word_count: int
    topics: List[str]
    sentiment: str
    elements: PageElements

# LangGraph state definitions
class CrawlState(TypedDict):
    current_url: str
    depth: int
    all_urls: List[Dict]  # {"url": str, "hasNewLinks": bool, "isVisited": bool}
    screenshot_paths: List[str]
    section_details: Dict  # {url: {"forms": [], "sections": [], "widgets": []}}
    visited: set
    to_crawl: List[tuple]  # (url, depth)
    analyses: List[Dict]  # Store PageAnalysis as dicts

class AnalysisState(TypedDict):
    content: str
    url: str
    summary: str
    topics: List[str]
    sentiment: str
    word_count: int
    metadata: Dict
    elements: Dict  # PageElements as dict

class AgnoBrowserCrawler:
    def __init__(self, url, section, sections, app_id, username="", password="", allUrls=[]):
        # Validate URL
        parsed_url = urlparse(url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise ValueError(f"Invalid URL: {url}")
        if parsed_url.scheme not in ["http", "https"]:
            raise ValueError(f"Unsupported URL scheme: {parsed_url.scheme}")
        if "https://" in url[8:]:
            logger.error(f"Malformed URL detected: {url}")
            raise ValueError(f"Malformed URL with duplicate protocol: {url}")

        self.url = url
        self.section = section
        self.sections = sections
        self.app_id = app_id
        self.username = username
        self.password = password
        self.screenshot_dir = "screenshots"
        self.screenshot_paths = []
        self.section_details = {}
        self.allUrls = allUrls

        if not os.path.exists(self.screenshot_dir):
            os.makedirs(self.screenshot_dir)

        logger.info(f"Initialized crawler for URL: {self.url}")

    async def create_page(self, context):
        try:
            page = await context.new_page()
            return page
        except Exception as e:
            logger.error(f"Error creating new page: {e}")
            raise

    async def launch_browser(self, context):
        page = await self.create_page(context)
        if self.username and self.password and self.username.strip() and self.password.strip():
            username_selector = None
            username_candidates = [
                'input[placeholder*="username" i]',
                'input[placeholder*="email" i]',
                'input[name*="username" i]',
                'input[name*="email" i]',
                'input[id*="username" i]',
                'input[id*="email" i]',
                'input[type="email" i]',
            ]
            for selector in username_candidates:
                element = await page.query_selector(selector)
                if element:
                    username_selector = selector
                    break
            if not username_selector:
                raise Exception("Could not find the username field")

            password_selector = None
            password_candidates = [
                'input[placeholder*="password" i]',
                'input[placeholder="Password" i]',
                'input[name*="password" i]',
                'input[id*="password" i]',
                'input[type="password" i]'
            ]
            for selector in password_candidates:
                element = await page.query_selector(selector)
                if element:
                    password_selector = selector
                    break
            if not password_selector:
                raise Exception("Could not find the password field")

            login_button_selector = None
            login_button_candidates = [
                'button[type="submit" i]',
                'button:has-text("login")',
                'button:has-text("sign in")',
                'button:has-text("log in")',
                'button:has-text("Login")',
                'button:has-text("Sign In")',
                'button:has-text("Log In")',
            ]
            for selector in login_button_candidates:
                element = await page.query_selector(selector)
                if element:
                    login_button_selector = selector
                    break
            if not login_button_selector:
                raise Exception("Could not find the login button")

            await page.fill(username_selector, self.username)
            await page.fill(password_selector, self.password)
            await page.click(login_button_selector)
            logger.info("Logged in successfully")
            wait_options = ["networkidle", "domcontentloaded", "load", "commit"]
            for wait_type in wait_options:
                try:
                    await page.wait_for_load_state(state=wait_type, timeout=TIMEOUT)
                    break
                except PlaywrightTimeoutError as e:
                    logger.warning(f"Failed waiting with {wait_type}: {e}")

        return page

    async def take_screenshot(self, page, name):
        try:
            screenshot_path = os.path.join(self.screenshot_dir, f"{self.app_id}_{name}.png")
            await page.screenshot(path=screenshot_path, full_page=True)
            logger.info(f"Screenshot saved: {screenshot_path}")
            return screenshot_path
        except PlaywrightError as e:
            logger.error(f"Error taking screenshot: {e}")
            return None

    async def close_popups(self, page):
        try:
            is_closed = await page.is_closed()
            if is_closed:
                raise PlaywrightError("Page is closed")
            close_buttons = await page.query_selector_all('button.close, [aria-label="Close"], [aria-label="Dismiss"]')
            for button in close_buttons:
                await button.click()
                await page.wait_for_timeout(1000)
            logger.info("Closed popups")
        except PlaywrightError as e:
            logger.info(f"No popups found or error closing popups: {e}")

    async def extract_links(self, crawler, url: str) -> List[str]:
        logger.info(f"Extracting links from {url}")
        schema = {
            "name": "Links",
            "baseSelector": "a",
            "fields": [
                {"name": "url", "selector": "a", "type": "attribute", "attribute": "href"},
                {"name": "text", "selector": "a", "type": "text"}
            ]
        }
        extraction = JsonCssExtractionStrategy(schema)
        run_conf = CrawlerRunConfig(
            extraction_strategy=extraction,
            cache_mode=CacheMode.BYPASS,
            exclude_external_links=True,
            exclude_social_media_links=True
        )
        result = await crawler.arun(url=url, config=run_conf)
        if not result.success:
            logger.error(f"Failed to extract links from {url}: {result.error_message}")
            return []

        extracted = json.loads(result.extracted_content)
        base_domain = urlparse(url).netloc
        links = []
        for item in extracted:
            link_url = item.get("url", "")
            if link_url:
                absolute_url = urljoin(url, link_url.strip())
                absolute_url, _ = urldefrag(absolute_url)
                if urlparse(absolute_url).netloc == base_domain and absolute_url.startswith("http"):
                    already_present = any(entry["url"] == absolute_url for entry in self.allUrls)
                    if not already_present and absolute_url not in links:
                        links.append(absolute_url)
        logger.info(f"Extracted {len(links)} unique internal links from {url}")
        return links

    async def scrape_and_extract(self, crawler, url: str) -> Dict:
        logger.info(f"Scraping and extracting elements from {url}")
        llm_config = LLMConfig(provider=f"ollama/{LLM_MODEL}", base_url=os.getenv("OLLAMA_BASE_URL", ""))
        extraction_strategy = LLMExtractionStrategy(
            llm_config=llm_config,
            schema=PageElements.model_json_schema(),
            extraction_type="schema",
            instruction="Extract all forms, sections, and widgets from the page content. "
                        "Forms should include action, method, inputs (type, name, placeholder), and whether it's a login form (has password field or login-related terms). "
                        "Sections include <section>, <nav>, or <div> with navigation roles or menu classes. "
                        "Widgets include <iframe>, sliders, modals, or carousels. Return valid JSON.",
            chunk_token_threshold=CHUNK_TOKEN_THRESHOLD,
            overlap_rate=OVERLAP_RATE,
            apply_chunking=True,
            input_format="markdown",
            extra_args={"temperature": 0.1, "max_tokens": 1500},
            verbose=True
        )
        run_conf = CrawlerRunConfig(extraction_strategy=extraction_strategy, cache_mode=CacheMode.BYPASS)
        result = await crawler.arun(url=url, config=run_conf)
        if not result.success:
            logger.error(f"Failed to extract elements from {url}: {result.error_message}")
            return {"url": url, "markdown": "", "metadata": {}, "elements": {"forms": [], "sections": [], "widgets": []}}

        try:
            elements = json.loads(result.extracted_content)
        except json.JSONDecodeError:
            logger.error(f"Invalid JSON from LLM for {url}")
            elements = {"forms": [], "sections": [], "widgets": []}

        logger.info(f"Extracted elements from {url}: {len(elements['forms'])} forms, {len(elements['sections'])} sections, {len(elements['widgets'])} widgets")
        return {
            "url": url,
            "markdown": result.markdown,
            "metadata": result.metadata,
            "elements": elements
        }

    def truncate_content(self, content: str, max_words: int = MAX_TOKENS_CONTENT) -> str:
        words = content.split()
        if len(words) > max_words:
            logger.info(f"Truncating content from {len(words)} to {max_words} words")
            return " ".join(words[:max_words])
        return content

    async def analyze_page_with_langgraph(self, content: str, url: str, scraped_data: Dict) -> PageAnalysis:
        logger.info(f"Analyzing page: {url}")
        try:
            llm = ChatOllama(model=LLM_MODEL, base_url=os.getenv("OLLAMA_BASE_URL", ""))

            async def summarize(state: AnalysisState) -> AnalysisState:
                prompt = f"Summarize the following content in 2-3 sentences:\n{state['content']}"
                response = await llm.ainvoke(prompt)
                state["summary"] = response.content if response.content else "N/A"
                logger.info(f"Summary for {state['url']}: {state['summary']}")
                return state

            async def identify_topics(state: AnalysisState) -> AnalysisState:
                prompt = f"List up to 5 main topics discussed in the following content:\n{state['content']}"
                response = await llm.ainvoke(prompt)
                topics = response.content.split("\n") if response.content else []
                state["topics"] = [t.strip() for t in topics if t.strip()]
                logger.info(f"Topics for {state['url']}: {state['topics']}")
                return state

            async def analyze_sentiment(state: AnalysisState) -> AnalysisState:
                prompt = f"Determine the sentiment (positive, negative, neutral) of the following content:\n{state['content']}"
                response = await llm.ainvoke(prompt)
                state["sentiment"] = response.content if response.content else "N/A"
                logger.info(f"Sentiment for {state['url']}: {state['sentiment']}")
                return state

            workflow = StateGraph(AnalysisState)
            workflow.add_node("summarize", summarize)
            workflow.add_node("identify_topics", identify_topics)
            workflow.add_node("analyze_sentiment", analyze_sentiment)
            workflow.add_edge("summarize", "identify_topics")
            workflow.add_edge("identify_topics", "analyze_sentiment")
            workflow.add_edge("analyze_sentiment", END)
            workflow.set_entry_point("summarize")
            graph = workflow.compile()

            initial_state = {
                "content": self.truncate_content(content),
                "url": url,
                "summary": "N/A",
                "topics": [],
                "sentiment": "N/A",
                "word_count": len(content.split()),
                "metadata": scraped_data.get("metadata", {}),
                "elements": scraped_data.get("elements", {"forms": [], "sections": [], "widgets": []})
            }

            final_state = await graph.ainvoke(initial_state)

            return PageAnalysis(
                url=final_state["url"],
                title=final_state["metadata"].get("title", "N/A"),
                summary=final_state["summary"],
                word_count=final_state["word_count"],
                topics=final_state["topics"],
                sentiment=final_state["sentiment"],
                elements=PageElements(**final_state["elements"])
            )
        except Exception as e:
            logger.error(f"LangGraph analysis failed for {url}: {e}")
            return PageAnalysis(
                url=url,
                title="N/A",
                summary="Failed to analyze",
                word_count=0,
                topics=[],
                sentiment="N/A",
                elements=PageElements(forms=[], sections=[], widgets=[])
            )

    def create_docx(self, analyses: List[PageAnalysis]):
        doc = Document()
        doc.add_heading("Website Analysis Report (Crawl4AI + LangGraph)", 0)
        doc.add_paragraph(f"Total URLs analyzed: {len(analyses)}")

        for analysis in analyses:
            doc.add_heading(analysis.url, level=1)
            doc.add_paragraph(f"**Title**: {analysis.title}")
            doc.add_paragraph(f"**Summary**: {analysis.summary}")
            doc.add_paragraph(f"**Word Count**: {analysis.word_count}")
            doc.add_paragraph(f"**Topics**: {', '.join(analysis.topics)}")
            doc.add_paragraph(f"**Sentiment**: {analysis.sentiment}")
            doc.add_paragraph(f"**Forms**: {len(analysis.elements.forms)} found")
            for form in analysis.elements.forms:
                doc.add_paragraph(f"  - Action: {form.action}, Method: {form.method}, Is Login: {form.is_login}")
            doc.add_paragraph(f"**Sections**: {len(analysis.elements.sections)} found")
            for section in analysis.elements.sections:
                doc.add_paragraph(f"  - Tag: {section.tag}, Class: {section.class_name}, ID: {section.id}")
            doc.add_paragraph(f"**Widgets**: {len(analysis.elements.widgets)} found")
            for widget in analysis.elements.widgets:
                doc.add_paragraph(f"  - Type: {widget.type}, Selector: {widget.selector}")
            doc.add_paragraph("\n")

        doc.save(OUTPUT_DOC)
        logger.info(f"Analysis saved to {OUTPUT_DOC}")
        print(f"\n✅ Analysis saved to {OUTPUT_DOC}")

    async def navigate_and_screenshot(self, state: CrawlState, crawler, context) -> CrawlState:
        url = state["current_url"]
        depth = state["depth"]
        logger.info(f"Crawling {url} at depth {depth}")

        if url in {u.strip() for u in os.getenv("CRAWLER_SKIP_URLS", "").split(",") if u.strip()}:
            logger.info(f"Skipping configured URL: {url}")
            state["visited"].add(url)
            return state

        # Use Playwright for navigation and screenshot
        page = await self.create_page(context)
        try:
            await page.goto(url, wait_until=WAIT_UNTIL, timeout=TIMEOUT)
            await self.close_popups(page)
            screenshot_name = f"{uuid.uuid4().hex[:8]}"
            screenshot_path = await self.take_screenshot(page, screenshot_name)
            if screenshot_path:
                state["screenshot_paths"].append(screenshot_path)

            # Scrape and extract elements with LLM
            scraped_data = await self.scrape_and_extract(crawler, url)
            analysis = await self.analyze_page_with_langgraph(scraped_data["markdown"], url, scraped_data)
            state["analyses"].append(analysis.dict())
            state["section_details"][url] = {
                "forms": [form.dict() for form in analysis.elements.forms],
                "sections": [section.dict() for section in analysis.elements.sections],
                "widgets": [widget.dict() for widget in analysis.elements.widgets]
            }
        except PlaywrightTimeoutError as e:
            logger.error(f"Navigation timeout for {url}: {e}")
        finally:
            await page.close()

        state["visited"].add(url)
        return state

    async def extract_and_queue_links(self, state: CrawlState, crawler, context) -> CrawlState:
        url = state["current_url"]
        depth = state["depth"]

        # Extract links using Crawl4AI
        links = await self.extract_links(crawler, url)
        new_links = []
        for link in links:
            if link not in [entry["url"] for entry in state["all_urls"]]:
                new_links.append(link)
                state["all_urls"].append({"url": link, "hasNewLinks": False, "isVisited": False})

        # Update current URL's metadata
        for entry in state["all_urls"]:
            if entry["url"].rstrip('/') == url.rstrip('/'):
                entry["hasNewLinks"] = len(new_links) > 1
                entry["isVisited"] = True
                break

        # Queue new links for crawling
        if depth < MAX_DEPTH:
            state["to_crawl"].extend((link, depth + 1) for link in new_links if link not in state["visited"])

        logger.info(f"Queued {len(new_links)} new links from {url}")
        return state

    async def select_next_url(self, state: CrawlState, crawler, context) -> CrawlState:
        while state["to_crawl"] and len(state["all_urls"]) < MAX_URLS:
            url, depth = state["to_crawl"].pop(0)
            if url not in state["visited"]:
                state["current_url"] = url
                state["depth"] = depth
                logger.info(f"Selected next URL: {url} at depth {depth}")
                return state
        logger.info("No more URLs to crawl")
        return state

    async def crawl(self):
        # Define the workflow graph
        workflow = StateGraph(CrawlState)

        # Define async node functions with explicit crawler and context
        async def navigate_and_screenshot_node(state: CrawlState, crawler, context) -> CrawlState:
            return await self.navigate_and_screenshot(state, crawler, context)

        async def extract_and_queue_links_node(state: CrawlState, crawler, context) -> CrawlState:
            return await self.extract_and_queue_links(state, crawler, context)

        async def select_next_url_node(state: CrawlState, crawler, context) -> CrawlState:
            return await self.select_next_url(state, crawler, context)

        # Add nodes with lambda to pass crawler and context
        workflow.add_node("navigate_and_screenshot", lambda state, config: navigate_and_screenshot_node(state, config["crawler"], config["context"]))
        workflow.add_node("extract_and_queue_links", lambda state, config: extract_and_queue_links_node(state, config["crawler"], config["context"]))
        workflow.add_node("select_next_url", lambda state, config: select_next_url_node(state, config["crawler"], config["context"]))

        # Define edges
        workflow.add_edge("navigate_and_screenshot", "extract_and_queue_links")
        workflow.add_edge("extract_and_queue_links", "select_next_url")
        workflow.add_conditional_edges(
            "select_next_url",
            lambda state: "navigate_and_screenshot" if state["to_crawl"] and len(state["all_urls"]) < MAX_URLS else END,
            {"navigate_and_screenshot": "navigate_and_screenshot", END: END}
        )
        workflow.set_entry_point("navigate_and_screenshot")

        # Initialize state
        initial_state = {
            "current_url": self.url,
            "depth": 0,
            "all_urls": self.allUrls,
            "screenshot_paths": self.screenshot_paths,
            "section_details": self.section_details,
            "visited": set(),
            "to_crawl": [(self.url, 0)],
            "analyses": []
        }

        # Run the crawler
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=False)
            try:
                context = await browser.new_context(
                    viewport={'width': 1460, 'height': 1080},
                    record_video_dir="video_output",
                    record_video_size={"width": 640, "height": 480}
                )
                page = await self.launch_browser(context)
                await page.close()

                browser_conf = BrowserConfig(headless=True, viewport_width=1280, viewport_height=720)
                with tqdm_asyncio(total=MAX_URLS, desc="Crawling pages") as pbar:
                    async with AsyncWebCrawler(config=browser_conf) as crawler:
                        logger.info(f"Starting deep crawl of {self.url} with Crawl4AI")
                        print(f"🔍 Starting deep crawl of {self.url} with Crawl4AI...")
                        graph = workflow.compile()
                        final_state = await graph.ainvoke(initial_state, config={"crawler": crawler, "context": context})
                        pbar.update(len(final_state["all_urls"]))

                    # Process sections (if provided)
                    for section in self.sections:
                        section_url = urljoin(self.url, section)
                        if section_url not in [entry["url"] for entry in final_state["all_urls"]]:
                            initial_state["current_url"] = section_url
                            initial_state["depth"] = 0
                            initial_state["to_crawl"] = [(section_url, 0)]
                            async with AsyncWebCrawler(config=browser_conf) as crawler:
                                logger.info(f"Crawling section: {section_url}")
                                graph = workflow.compile()
                                final_state = await graph.ainvoke(initial_state, config={"crawler": crawler, "context": context})
                                pbar.update(len(final_state["all_urls"]))

                # Generate .docx report
                self.create_docx([PageAnalysis(**analysis) for analysis in final_state["analyses"]])

                # Prepare JSON output
                output = {
                    "screenshot_paths": final_state["screenshot_paths"],
                    "section_details": final_state["section_details"],
                    "allLink": final_state["all_urls"]
                }

                print(json.dumps(output, indent=3))
                logger.info("Crawl completed successfully")
                return output
            finally:
                await context.close()
                await browser.close()

async def main():
    data = {
        'url': os.getenv("POC_START_URL", ""),
        'section': 'forms',
        'sections': [],
        'app_id': os.getenv("POC_APP_ID", ""),
        'username': '',
        'password': ''
    }

    crawler = AgnoBrowserCrawler(
        url=data['url'],
        section=data['section'],
        allUrls=[],
        sections=data['sections'],
        app_id=data['app_id'],
        username=data['username'],
        password=data['password']
    )

    result = await crawler.crawl()

if __name__ == "__main__":
    asyncio.run(main())