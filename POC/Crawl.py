import os
import re
import asyncio
import argparse
import markdown2
from docx import Document
from docx.shared import Inches
import pdfkit
from datetime import datetime
from urllib.parse import urlparse
from playwright.async_api import async_playwright

# Import Agno and crawl4ai
from agno.agent import Agent
from agno.tools.crawl4ai import Crawl4aiTools

# For local summarization without API key dependency
try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False
    print("⚠️ Ollama not available. Will use simple summarization instead.")

# Default configs
DEFAULT_URL = os.getenv("POC_START_URL", "")
SCREENSHOT_DIR = "screenshots"
DEFAULT_OLLAMA_MODEL = "gemma:2b"  # A lightweight model for summaries
OUTPUT_DIR = "documentation"

# These variables can be modified by command line arguments
#global SCREENSHOT_DIR, OUTPUT_DIR

# Create necessary directories
os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

def url_to_filename(url: str) -> str:
    """Convert URL to a valid filename for screenshots."""
    slug = re.sub(r'https?://|[^\w\-]', '_', url)
    return f"{SCREENSHOT_DIR}/{slug[:80]}.png"

async def capture_screenshot(url: str, filepath: str):
    """Capture a full-page screenshot of the given URL."""
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            # Increase timeout and add navigation options for more reliable capture
            await page.goto(url, timeout=30000, wait_until="networkidle")
            await page.screenshot(path=filepath, full_page=True)
            await browser.close()
            return True
    except Exception as e:
        print(f"❌ Screenshot error for {url}: {e}")
        return False

def get_domain_name(url):
    """Extract domain name from URL for file naming."""
    parsed_url = urlparse(url)
    domain = parsed_url.netloc
    return domain.replace("www.", "")

def simple_summarize(url, text):
    """Fallback summarization when no LLM is available."""
    # Extract title (if any)
    title_match = re.search(r'<title>(.*?)</title>', text, re.IGNORECASE)
    title = title_match.group(1) if title_match else "Untitled Page"

    # Count forms, buttons, and links
    forms_count = len(re.findall(r'<form', text, re.IGNORECASE))
    buttons_count = len(re.findall(r'<button|type="button"|type="submit"', text, re.IGNORECASE))
    links_count = len(re.findall(r'<a\s+', text, re.IGNORECASE))

    # Extract meta description (if any)
    meta_desc_match = re.search(r'<meta\s+name="description"\s+content="(.*?)"', text, re.IGNORECASE)
    meta_desc = meta_desc_match.group(1) if meta_desc_match else "No description available"

    # Create a simple summary
    summary = f"""
# {title}

## Page Overview
- URL: {url}
- Description: {meta_desc}

## Interactive Elements
- Forms: {forms_count}
- Buttons: {buttons_count}
- Links: {links_count}

## Content Snippet
```
{text[:300]}...
```
    """
    return summary

async def run_agno_crawler(url, max_pages=None):
    """Run the Agno agent with crawl4ai tools to crawl the website."""
    print(f"🚀 Starting Agno Agent to crawl {url}...")

    crawl4ai_tools = Crawl4aiTools(max_length=None)

    agent = Agent(
        name="WebsiteCrawler",
        description=f"Crawls and documents {url} with comprehensive analysis.",
        tools=[crawl4ai_tools],
        show_tool_calls=True
    )

    # Run the agent to crawl the website
    response = agent.run(f"Scrape all content from {url} including forms, login pages, and interactive elements.")

    # Parse the response based on the agent's output format
    pages = []
    if hasattr(response, 'content') and response.content:
        if isinstance(response.content, dict):
            pages = response.content.get("pages", [])
        elif isinstance(response.content, list):
            pages = response.content
    else:
        print("⚠️ No content returned from agent")

    return pages

async def generate_documentation(url, model=DEFAULT_OLLAMA_MODEL, max_pages=None):
    """Generate comprehensive documentation for the crawled website."""
    domain_name = get_domain_name(url)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")

    # Define output files
    output_md = f"{OUTPUT_DIR}/{domain_name}_{timestamp}_summary.md"
    output_pdf = f"{OUTPUT_DIR}/{domain_name}_{timestamp}_summary.pdf"
    output_docx = f"{OUTPUT_DIR}/{domain_name}_{timestamp}_summary.docx"

    # Crawl the website using Agno
    pages = await run_agno_crawler(url, max_pages)

    if not pages:
        print("❌ No pages were crawled. Check the URL and try again.")
        return

    print(f"✅ Crawled {len(pages)} pages")

    # Process each page to get screenshots and summaries
    summaries = {}

    for page in pages:
        page_url = page.get("url", "unknown")
        text = page.get("text", "").strip()
        print(f"\n📄 Processing {page_url}")

        # Take a screenshot of the page
        screenshot_path = url_to_filename(page_url)
        screenshot_success = await capture_screenshot(page_url, screenshot_path)

        if not text:
            summaries[page_url] = ("❌ No extractable text found.", screenshot_path if screenshot_success else None)
            continue

        # Generate a summary of the page
        if OLLAMA_AVAILABLE:
            try:
                prompt = f"""
                Summarize this webpage: {page_url}
                Highlight if it contains:
                - Forms (login, registration, input fields)
                - Buttons or interactive elements
                - Sections with demos or visual UI components

                Content:
                {text[:2000]}
                """
                res = ollama.chat(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You're a web documentation summarizer."},
                        {"role": "user", "content": prompt}
                    ]
                )
                summary = res['message']['content']
            except Exception as e:
                print(f"⚠️ Ollama error: {e}. Falling back to simple summarization.")
                summary = simple_summarize(page_url, text)
        else:
            summary = simple_summarize(page_url, text)

        summaries[page_url] = (summary, screenshot_path if screenshot_success else None)

    # Write Markdown documentation
    print(f"\n📝 Writing Markdown to {output_md}")
    with open(output_md, "w", encoding="utf-8") as f:
        f.write(f"# {domain_name.upper()} Website Documentation\n\n")
        f.write(f"Crawl of {url} completed on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        f.write(f"Total pages crawled: {len(pages)}\n\n")

        # Table of contents
        f.write("## Table of Contents\n\n")
        for i, page_url in enumerate(summaries.keys()):
            parsed = urlparse(page_url)
            page_label = parsed.path if parsed.path else "Homepage"
            f.write(f"{i+1}. [{page_label}](#{i+1})\n")

        f.write("\n---\n\n")

        # Page details
        for i, (page_url, (summary, img_path)) in enumerate(summaries.items()):
            parsed = urlparse(page_url)
            page_label = parsed.path if parsed.path else "Homepage"

            f.write(f"## {i+1}. {page_label}\n\n")
            f.write(f"**URL:** {page_url}\n\n")

            if img_path and os.path.exists(img_path):
                img_rel = os.path.relpath(img_path, os.path.dirname(output_md))
                img_rel = img_rel.replace("\\", "/")
                f.write(f"![Screenshot]({img_rel})\n\n")

            f.write(f"{summary}\n\n---\n\n")

    # Generate PDF
    print(f"🧾 Exporting PDF to {output_pdf}")
    try:
        markdown_html = markdown2.markdown_path(output_md)
        pdfkit.from_string(markdown_html, output_pdf)
    except Exception as e:
        print(f"❌ PDF generation failed: {e}")
        print("💡 Note: pdfkit requires wkhtmltopdf to be installed")

    # Generate DOCX
    print(f"🧾 Exporting DOCX to {output_docx}")
    try:
        doc = Document()
        doc.add_heading(f"{domain_name.upper()} Website Documentation", 0)
        doc.add_paragraph(f"Crawl of {url} completed on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total pages crawled: {len(pages)}")

        # Add table of contents
        doc.add_heading("Table of Contents", level=1)
        for i, page_url in enumerate(summaries.keys()):
            parsed = urlparse(page_url)
            page_label = parsed.path if parsed.path else "Homepage"
            doc.add_paragraph(f"{i+1}. {page_label}")

        doc.add_page_break()

        # Add page details
        for i, (page_url, (summary, img_path)) in enumerate(summaries.items()):
            parsed = urlparse(page_url)
            page_label = parsed.path if parsed.path else "Homepage"

            doc.add_heading(f"{i+1}. {page_label}", level=1)
            doc.add_paragraph(f"URL: {page_url}")

            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                except Exception as e:
                    doc.add_paragraph(f"[Image could not be added: {e}]")

            doc.add_paragraph(summary)
            doc.add_page_break()

        doc.save(output_docx)
    except Exception as e:
        print(f"❌ DOCX generation failed: {e}")

    print("\n✅ Done! Outputs:")
    print(f"- Markdown: {output_md}")
    print(f"- PDF:      {output_pdf}")
    print(f"- DOCX:     {output_docx}")

    return {
        "markdown": output_md,
        "pdf": output_pdf,
        "docx": output_docx
    }

async def main():
    global OUTPUT_DIR, SCREENSHOT_DIR

    # Parse command line arguments
    parser = argparse.ArgumentParser(description="Website Documentation Generator with Agno and crawl4ai")
    parser.add_argument("--url", default=DEFAULT_URL, help=f"URL to crawl (default: {DEFAULT_URL})")
    parser.add_argument("--model", default=DEFAULT_OLLAMA_MODEL, help=f"Ollama model for summaries (default: {DEFAULT_OLLAMA_MODEL})")
    parser.add_argument("--output-dir", default=OUTPUT_DIR, help=f"Output directory (default: {OUTPUT_DIR})")
    parser.add_argument("--max-pages", type=int, default=None, help="Maximum pages to crawl (default: no limit)")

    args = parser.parse_args()

    # Update global variables based on args
    OUTPUT_DIR = args.output_dir
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    try:
        await generate_documentation(args.url, args.model, args.max_pages)
    except KeyboardInterrupt:
        print("\n⚠️ Process interrupted by user")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    asyncio.run(main())