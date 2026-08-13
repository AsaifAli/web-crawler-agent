import os
import uuid
import json
import asyncio
import logging
import shutil
import time
import random
from datetime import datetime
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
import markdown2
import easyocr
from transformers import pipeline
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError, Error as PlaywrightError
from urllib.parse import urljoin, urlparse
from crawl4ai import AsyncWebCrawler, CrawlerRunConfig, CacheMode
from functools import wraps
from http.client import HTTPException

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def retry_with_backoff(max_retries=1, base_delay=5, max_delay=60):
    def decorator(func):
        @wraps(func)
        async def wrapper(*args, **kwargs):
            for attempt in range(max_retries):
                try:
                    return await func(*args, **kwargs)
                except Exception as e:
                    if attempt == max_retries - 1:
                        logger.error(f"{func.__name__} failed after {max_retries} attempts: {e}", exc_info=True)
                        raise
                    delay = min(base_delay * (2 ** attempt) + random.uniform(0, 0.1), max_delay)
                    logger.warning(f"{func.__name__} attempt {attempt + 1}/{max_retries} failed with {e}. Retrying in {delay:.2f}s...")
                    await asyncio.sleep(delay)
        return wrapper
    return decorator

class ImageContent(BaseModel):
    filename: str
    timestamp: str
    extracted_text: str
    summary: str

class ImageToDocProcessor:
    def __init__(self, image_folder="screenshots", app_id="webcrawler"):
        self.image_folder = image_folder
        self.app_id = app_id
        self.image_details = {}
        self.report_dir = "reports"
        self.debug_dir = "debug"
        try:
            self.reader = easyocr.Reader(['en'], gpu=False)  # Initialize EasyOCR for English
            logger.info("EasyOCR initialized successfully.")
        except Exception as e:
            logger.error(f"Error initializing EasyOCR: {e}", exc_info=True)
            self.reader = None
        try:
            self.summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-6-6")
            logger.info("DistilBART summarization model initialized successfully.")
        except Exception as e:
            logger.error(f"Error initializing DistilBART summarization model: {e}", exc_info=True)
            self.summarizer = None
        for directory in [self.image_folder, self.report_dir, self.debug_dir]:
            os.makedirs(directory, exist_ok=True)

    @retry_with_backoff(max_retries=1, base_delay=5, max_delay=60)
    async def extract_text_from_image(self, image_path: str) -> str:
        if not os.path.exists(image_path):
            logger.error(f"Image file not found for OCR processing: {image_path}")
            return ""
        if not self.reader:
            logger.error("EasyOCR not initialized, skipping text extraction.")
            return ""
        try:
            logger.debug(f"Processing image with EasyOCR: {image_path}")
            results = self.reader.readtext(image_path, detail=0, paragraph=True)
            extracted_text = " ".join(results).strip()
            logger.info(f"EasyOCR extracted text from {image_path}: {len(extracted_text)} characters")
            debug_text_path = os.path.join(self.debug_dir, f"ocr_extracted_text_{os.path.basename(image_path)}.txt")
            with open(debug_text_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
            await asyncio.sleep(1)  # Small delay to simulate async processing
            return extracted_text
        except Exception as e:
            logger.error(f"Error extracting text using EasyOCR from {image_path}: {e}", exc_info=True)
            return ""

    @retry_with_backoff(max_retries=1, base_delay=5, max_delay=60)
    async def summarize_content(self, content: str) -> str:
        if not content:
            logger.warning("No content provided for summarization.")
            return ""
        if not self.summarizer:
            logger.warning("Summarization skipped: DistilBART model not available.")
            return content[:200] + "..." if content else ""
        logger.info("Generating summary with DistilBART model")
        try:
            max_summary_chars = 4000
            truncated_content = content[:max_summary_chars]
            if len(content) > max_summary_chars:
                logger.warning(f"Content truncated to {max_summary_chars} chars for summarization.")
            summary = self.summarizer(
                truncated_content,
                max_length=120,
                min_length=40,
                do_sample=False
            )[0]['summary_text'].strip()
            logger.info("DistilBART model generated summary successfully")
            summary_debug_path = os.path.join(self.debug_dir, f"distilbart_summary_{uuid.uuid4().hex[:8]}.txt")
            with open(summary_debug_path, 'w', encoding='utf-8') as f:
                f.write(f"---CONTENT---\n{truncated_content}\n\n---SUMMARY---\n{summary}")
            await asyncio.sleep(1)  # Small delay to simulate async processing
            return summary
        except Exception as e:
            logger.error(f"Error summarizing content with DistilBART model: {e}", exc_info=True)
            return content[:200] + "..." if content else ""

    async def scan_images(self):
        logger.info(f"Scanning images in {self.image_folder} using OCR Processor")
        try:
            image_files = [f for f in os.listdir(self.image_folder) if f.lower().endswith('.png')]
        except FileNotFoundError:
            logger.error(f"Image folder not found: {self.image_folder}")
            return
        except Exception as e:
            logger.error(f"Error listing files in {self.image_folder}: {e}")
            return
        if not image_files:
            logger.warning(f"No .png images found in {self.image_folder}")
            return
        for image_file in image_files:
            image_path = os.path.join(self.image_folder, image_file)
            if not os.path.isfile(image_path):
                logger.warning(f"Skipping non-file item: {image_path}")
                continue
            try:
                try:
                    parts = os.path.splitext(image_file)[0].split('_')
                    timestamp_str = next((p for p in reversed(parts) if len(p) >= 15 and p[8] == '_'), None)
                    timestamp = timestamp_str or datetime.fromtimestamp(os.path.getmtime(image_path)).strftime("%Y%m%d_%H%M%S")
                except Exception:
                    timestamp = datetime.fromtimestamp(os.path.getmtime(image_path)).strftime("%Y%m%d_%H%M%S")
                extracted_text = await self.extract_text_from_image(image_path)
                summary = await self.summarize_content(extracted_text) if extracted_text else ""
                self.image_details[image_file] = {
                    "filename": image_file,
                    "timestamp": timestamp,
                    "extracted_text": extracted_text,
                    "summary": summary,
                    "image_path": image_path
                }
                logger.info(f"Processed image via OCR: {image_file}")
            except Exception as e:
                logger.error(f"Error processing image {image_file} with OCR Processor: {e}", exc_info=True)

    def generate_report(self):
        logger.info("Generating reports based on OCR processing results...")
        if not self.image_details:
            logger.warning("No image details (OCR), skipping reports.")
            return None, None, None
        report_base_name = f"{self.app_id}_ocr_image_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        md_content = f"# Image Processing Report (OCR)\n\n"
        md_content += f"**Generated on**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        md_content += f"**Application ID**: {self.app_id}\n"
        md_content += f"**Image Folder**: {self.image_folder}\n\n## Processed Images\n\n"
        sorted_items = sorted(self.image_details.items(), key=lambda item: item[1]['timestamp'])
        doc = Document()
        doc.add_heading("Image Processing Report (OCR)", 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Application ID: {self.app_id}")
        doc.add_heading("Processed Images", level=1)
        for image_file, details in sorted_items:
            md_content += f"### Image: `{image_file}`\n"
            md_content += f"- **Timestamp**: {details['timestamp']}\n"
            md_content += f"- **Summary (DistilBART)**: {details['summary'] or 'N/A'}\n"
            relative_image_path = os.path.join(os.path.basename(self.image_folder), image_file)
            md_content += f"- **Image**: ![{image_file}]({relative_image_path})\n"
            text_preview = (details['extracted_text'][:200].replace('\n', ' ') + '...' if len(details['extracted_text']) > 200 else details['extracted_text']) if details['extracted_text'] else 'N/A'
            md_content += f"- **Extracted Text (OCR Preview)**: {text_preview} ({len(details.get('extracted_text',''))} chars)\n\n"
            doc.add_heading(f"Image: {image_file}", level=2)
            doc.add_paragraph(f"Timestamp: {details['timestamp']}")
            doc.add_paragraph(f"Summary (DistilBART): {details['summary'] or 'N/A'}")
            text_preview_docx = (details['extracted_text'][:500].replace('\r\n', '\n') + '...' if len(details['extracted_text']) > 500 else details['extracted_text']) if details['extracted_text'] else 'N/A'
            doc.add_paragraph(f"Extracted Text (OCR Preview): {text_preview_docx} ({len(details.get('extracted_text',''))} characters)")
            abs_image_path = os.path.abspath(details['image_path'])
            if os.path.exists(abs_image_path):
                try:
                    doc.add_picture(abs_image_path, width=Inches(6.0))
                except Exception as e:
                    logger.error(f"Error adding image {abs_image_path} to DOCX: {e}")
                    doc.add_paragraph(f"[Error adding image: {e}]")
            else:
                logger.warning(f"Image not found for DOCX: {abs_image_path}")
                doc.add_paragraph(f"[Image file not found: {image_file}]")
            doc.add_paragraph()
        md_path = os.path.join(self.report_dir, f"{report_base_name}.md")
        try:
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            logger.info(f"Markdown report (OCR) saved: {md_path}")
        except Exception as e:
            logger.error(f"Error saving MD report (OCR): {e}", exc_info=True)
            md_path = None
        html_path = os.path.join(self.report_dir, f"{report_base_name}.html")
        try:
            html_body = markdown2.markdown(md_content, extras=["tables", "fenced-code-blocks"])
            html_full = f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>Image Processing Report (OCR) - {self.app_id}</title><style>body{{font-family:sans-serif;margin:20px;line-height:1.6}}h1,h2,h3{{color:#333}}code{{background-color:#f0f0f0;padding:2px 4px;border-radius:3px}}img{{max-width:600px;height:auto;display:block;margin-top:10px;border:1px solid #ddd}}ul{{padding-left:20px}}li{{margin-bottom:10px}}</style></head><body>{html_body}</body></html>"""
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_full)
            logger.info(f"HTML Report (OCR) saved: {html_path}")
        except Exception as e:
            logger.error(f"Error saving HTML report (OCR): {e}", exc_info=True)
            html_path = None
        docx_path = os.path.join(self.report_dir, f"{report_base_name}.docx")
        try:
            doc.save(docx_path)
            logger.info(f"DOCX report (OCR) saved: {docx_path}")
        except Exception as e:
            logger.error(f"Error saving DOCX report (OCR): {e}", exc_info=True)
            docx_path = None
        return md_path, html_path, docx_path

    async def process(self):
        logger.info("Running ImageToDocProcessor (OCR) for image scanning and document creation")
        start_time = datetime.now()
        try:
            await self.scan_images()
            if not self.image_details:
                logger.warning("No images processed (OCR), skipping reports")
                return {
                    "processor": "OCR",
                    "processed_image_count": 0,
                    "image_details": {},
                    "report_path": None,
                    "html_report_path": None,
                    "docx_report_path": None
                }
            md_path, html_path, docx_path = self.generate_report()
            result = {
                "processor": "OCR",
                "processed_image_count": len(self.image_details),
                "image_details": {
                    fname: {
                        "filename": d["filename"],
                        "timestamp": d["timestamp"],
                        "summary": d["summary"],
                        "extracted_text_length": len(d.get("extracted_text", "")),
                        "image_path": d["image_path"]
                    } for fname, d in self.image_details.items()
                },
                "report_path": md_path,
                "html_report_path": html_path,
                "docx_report_path": docx_path
            }
            json_summary_path = os.path.join(self.report_dir, f"{self.app_id}_ocr_processing_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
            try:
                with open(json_summary_path, 'w', encoding='utf-8') as f:
                    json.dump(result, f, indent=2)
                logger.info(f"Processing summary (OCR) saved: {json_summary_path}")
            except Exception as e:
                logger.error(f"Error saving processing summary JSON (OCR): {e}")
            duration = datetime.now() - start_time
            logger.info("Image processing (OCR) completed successfully!")
            logger.info(f"Images processed: {len(self.image_details)}")
            logger.info(f"Reports: MD: {md_path or 'F'}, HTML: {html_path or 'F'}, DOCX: {docx_path or 'F'}")
            logger.info(f"Processing duration: {duration}")
            return result
        except Exception as e:
            logger.error(f"Error during image processing (OCR): {e}", exc_info=True)
            return {
                "processor": "OCR",
                "processed_image_count": 0,
                "error": str(e),
                "image_details": {},
                "report_path": None,
                "html_report_path": None,
                "docx_report_path": None
            }

class AgnoBrowserCrawler:
    def __init__(self, url, section, sections, app_id, username="", password="", allUrls=None):
        self.url = url
        self.section = section
        self.sections = sections
        self.app_id = app_id
        self.username = username
        self.password = password
        self.screenshot_dir = "screenshots"
        self.screenshot_paths = []
        self.section_details = {}
        self.visited_urls = set()
        self.debug_dir = "debug"
        self.report_dir = "reports"
        self.image_processor = None
        for directory in [self.screenshot_dir, self.debug_dir, self.report_dir]:
            os.makedirs(directory, exist_ok=True)

    async def create_page(self, context):
        try:
            page = await context.new_page()
            logger.info("Created new Playwright page")
            return page
        except Exception as e:
            logger.error(f"Error creating new page: {e}", exc_info=True)
            raise

    async def attempt_login(self, page: 'Page'):
        if not self.username or not self.password:
            logger.info("No credentials provided, skipping login.")
            return False
        logger.info(f"Attempting login for URL: {page.url}")
        username_sel = 'input[name="username"], input[id="username"], input[type="email"]'
        password_sel = 'input[name="password"], input[id="password"], input[type="password"]'
        button_sel = 'button[type="submit"], button:has-text("Log in"), button:has-text("Sign in"), input[type="submit"]'
        try:
            username_field = await page.wait_for_selector(username_sel, state="visible", timeout=5000)
            password_field = await page.wait_for_selector(password_sel, state="visible", timeout=5000)
            login_button = await page.wait_for_selector(button_sel, state="visible", timeout=5000)
            if username_field and password_field and login_button:
                logger.info("Found login elements. Filling form.")
                await username_field.fill(self.username)
                await password_field.fill(self.password)
                await page.wait_for_timeout(500)
                await login_button.click()
                try:
                    success_selector = '[id="flash"].success, a[href="/logout"], div#flash:contains("You logged into a secure area!")'
                    await page.wait_for_selector(success_selector, state="visible", timeout=10000)
                    logger.info("Login successful: Verified secure page content.")
                    await page.wait_for_load_state("networkidle", timeout=20000)
                except PlaywrightTimeoutError:
                    logger.warning("Could not verify login success, proceeding.")
                return True
            else:
                logger.warning("Could not find all login elements.")
                return False
        except PlaywrightTimeoutError:
            logger.warning("Timeout finding login elements, skipping form login.")
            return False
        except Exception as e:
            logger.error(f"Error during login attempt: {e}", exc_info=True)
            return False

    async def take_screenshot(self, page, name_suffix):
        try:
            safe_suffix = "".join(c if c.isalnum() else "_" for c in name_suffix)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            screenshot_name = f"{self.app_id}_{safe_suffix}_{timestamp}.png"
            screenshot_path = os.path.join(self.screenshot_dir, screenshot_name)
            await page.wait_for_load_state('domcontentloaded', timeout=15000)
            await page.wait_for_timeout(1000)
            await page.screenshot(path=screenshot_path, full_page=True)
            self.screenshot_paths.append(screenshot_path)
            logger.info(f"Screenshot saved: {screenshot_path}")
            return screenshot_path
        except Exception as e:
            logger.error(f"Error taking screenshot '{name_suffix}': {e}", exc_info=True)
            return None

    async def close_popups(self, page):
        if page.is_closed():
            logger.warning("Attempted to close popups on closed page.")
            return
        logger.debug(f"Checking for popups on {page.url}")
        popup_selectors = [
            '[aria-label*="close" i]', '[aria-label*="dismiss" i]', 'button:has-text("Accept")',
            'button:has-text("Agree")', 'button:has-text("OK")', 'button:has-text("Got it")',
            'button.close', 'button.modal-close', '[class*="cookie"] button', '[id*="cookie"] button',
            'div[role="dialog"] button[class*="close"]'
        ]
        closed_popup = False
        for selector in popup_selectors:
            try:
                elements = await page.query_selector_all(selector)
                for element in elements:
                    if await element.is_visible() and await element.is_enabled():
                        logger.info(f"Found popup '{selector}'. Attempting to click.")
                        try:
                            await element.click(timeout=5000)
                            closed_popup = True
                            await page.wait_for_timeout(1000)
                        except Exception as e:
                            logger.warning(f"Could not click popup '{selector}': {e}")
            except Exception as e:
                logger.warning(f"Error querying selector '{selector}': {e}")
        if closed_popup:
            logger.info("Popups closed successfully.")
        else:
            logger.debug("No actionable popups found.")

    async def extract_page_links_playwright(self, page: 'Page', base_url: str) -> list:
        links_found = []
        if page.is_closed():
            logger.warning("Attempting to extract links from closed page.")
            return links_found
        logger.debug(f"Playwright: Extracting links from: {page.url}")
        try:
            await page.wait_for_load_state('domcontentloaded', timeout=15000)
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(1500)
            link_elements = await page.query_selector_all(
                'a[href], button[onclick], [role="link"], [data-href], [data-url], .nav-link, .menu-item'
            )
            logger.info(f"Playwright: Found {len(link_elements)} potential link elements.")
            current_page_url_parsed = urlparse(page.url)
            for element in link_elements:
                try:
                    href = await element.get_attribute('href') or await element.get_attribute('data-href') or await element.get_attribute('data-url')
                    text = (await element.text_content() or "").strip()
                    if href and not href.startswith(('javascript:', '#', 'mailto:', 'tel:')):
                        absolute_url = urljoin(base_url, href)
                        parsed_absolute_url = urlparse(absolute_url)
                        if parsed_absolute_url.netloc == current_page_url_parsed.netloc:
                            normalized_url = parsed_absolute_url._replace(fragment="", query="").geturl().rstrip('/')
                            if normalized_url:
                                links_found.append({"url": normalized_url, "text": text or "Link", "source": "playwright"})
                        else:
                            logger.debug(f"Playwright: Skipping external link: {absolute_url}")
                except Exception as e:
                    logger.warning(f"Playwright: Error parsing link: {e}")
            unique_links = list({link['url']: link for link in links_found}.values())
            logger.info(f"Playwright: Extracted {len(unique_links)} unique, same-domain links.")
            return unique_links
        except Exception as e:
            logger.error(f"Playwright: Error extracting links: {e}", exc_info=True)
            return []

    async def extract_links_with_crawl4ai(self, url_to_extract: str) -> list:
        logger.info(f"Crawl4AI: Attempting link extraction for: {url_to_extract}")
        extracted_links = []
        try:
            async with AsyncWebCrawler() as crawler:
                run_config = CrawlerRunConfig(
                    cache_mode=CacheMode.BYPASS,
                    exclude_external_links=True,
                    exclude_social_media_links=True
                )
                result = await crawler.arun(url=url_to_extract, config=run_config)
                if result and result.success and result.links:
                    all_crawl4ai_links = result.links.get("internal", [])
                    logger.info(f"Crawl4AI: Found {len(all_crawl4ai_links)} raw internal links.")
                    current_page_url_parsed = urlparse(url_to_extract)
                    base_domain = current_page_url_parsed.netloc
                    for link_data in all_crawl4ai_links:
                        href = link_data.get('href')
                        if href:
                            try:
                                absolute_url = urljoin(url_to_extract, href)
                                parsed_absolute_url = urlparse(absolute_url)
                                if parsed_absolute_url.netloc == base_domain:
                                    normalized_url = parsed_absolute_url._replace(fragment="", query="").geturl().rstrip('/')
                                    if normalized_url:
                                        extracted_links.append({
                                            "url": normalized_url,
                                            "text": link_data.get('text', 'Link'),
                                            "source": "crawl4ai"
                                        })
                                else:
                                    logger.debug(f"Crawl4AI: Skipping external link: {absolute_url}")
                            except Exception as e:
                                logger.warning(f"Crawl4AI: Error processing link '{href}': {e}")
                        else:
                            logger.debug(f"Crawl4AI: Skipping link data with no href: {link_data}")
                elif result and not result.success:
                    logger.error(f"Crawl4AI: Failed crawl {url_to_extract}. Error: {result.error_message}")
                else:
                    logger.info(f"Crawl4AI: No links or result unsuccessful for {url_to_extract}")
        except Exception as e:
            logger.error(f"Crawl4AI: Error during link extraction {url_to_extract}: {e}", exc_info=True)
        unique_links = list({link['url']: link for link in extracted_links}.values())
        logger.info(f"Crawl4AI: Extracted {len(unique_links)} unique, same-domain links.")
        return unique_links

    async def crawl_url(self, context, page: 'Page', url_to_crawl: str, base_domain: str):
        normalized_url = url_to_crawl.rstrip('/')
        if normalized_url in self.visited_urls:
            logger.info(f"Skipping already visited URL: {normalized_url}")
            return page, []
        logger.info(f"Crawling URL: {url_to_crawl}")
        screenshot_path, combined_links = None, []
        try:
            target_goto_url = url_to_crawl
            parsed_u_for_auth = urlparse(url_to_crawl)
            if "/basic_auth" in parsed_u_for_auth.path.lower():
                logger.info("Applying Basic Auth credentials for /basic_auth URL")
                if "@" not in parsed_u_for_auth.netloc:
                    target_goto_url = url_to_crawl
                    if parsed_u_for_auth.query:
                        target_goto_url += f"?{parsed_u_for_auth.query}"
                    logger.debug(f"Modified URL for Basic Auth: {target_goto_url}")
                else:
                    logger.warning("URL may already contain auth info, using original.")
            await page.goto(target_goto_url, wait_until="networkidle", timeout=60000)
            final_url = page.url.rstrip('/')
            if final_url != normalized_url and urlparse(final_url).path != urlparse(normalized_url).path:
                logger.info(f"Redirected from {normalized_url} to {final_url}")
                if final_url in self.visited_urls:
                    logger.info(f"Skipping visited redirected URL: {final_url}")
                    return page, []
                url_to_crawl = final_url
                normalized_url = final_url
            self.visited_urls.add(normalized_url)
            await self.close_popups(page)
            url_path_part = urlparse(normalized_url).path.replace('/', '_').strip('_') or "homepage"
            screenshot_name_suffix = f"{url_path_part}_{uuid.uuid4().hex[:6]}"
            screenshot_path = await self.take_screenshot(page, screenshot_name_suffix)
            page_base_url = urlparse(page.url)._replace(path="", query="", fragment="").geturl()
            playwright_links = await self.extract_page_links_playwright(page, page_base_url)
            crawl4ai_links = await self.extract_links_with_crawl4ai(page.url)
            all_links_dict = {link['url']: link for link in playwright_links + crawl4ai_links}
            combined_links = list(all_links_dict.values())
            logger.info(f"Combined links (Playwright: {len(playwright_links)}, Crawl4AI: {len(crawl4ai_links)}): Total {len(combined_links)}")
            self.section_details[normalized_url] = {
                "url": normalized_url,
                "screenshot": screenshot_path or "Failed",
                "extracted_link_count": len(combined_links),
                "link_sources": {"playwright": len(playwright_links), "crawl4ai": len(crawl4ai_links)}
            }
            logger.info(f"Successfully processed URL: {normalized_url}")
        except Exception as e:
            logger.error(f"Error crawling URL {url_to_crawl}: {e}", exc_info=True)
            self.visited_urls.add(normalized_url)
            self.section_details[normalized_url] = {
                "url": normalized_url,
                "error": f"Crawling Error: {e}",
                "screenshot": None
            }
        return page, combined_links

    async def crawl(self, max_pages=10):
        logger.info(f"Starting crawl for base URL: {self.url}, Max Pages: {max_pages}")
        start_time = datetime.now()
        urls_to_visit = []
        initial_urls = [self.url] + [urljoin(self.url, s) for s in self.sections]
        for u in initial_urls:
            normalized = u.rstrip('/')
            if normalized not in self.visited_urls:
                urls_to_visit.append(normalized)
        pages_visited_count = 0
        base_domain = urlparse(self.url).netloc
        # Clear screenshots directory
        if os.path.exists(self.screenshot_dir):
            try:
                shutil.rmtree(self.screenshot_dir)
                logger.info(f"Cleared existing screenshots directory.")
            except Exception as e:
                logger.error(f"Error clearing screenshot directory: {e}")
        os.makedirs(self.screenshot_dir, exist_ok=True)
        # Clear reports directory
        if os.path.exists(self.report_dir):
            try:
                shutil.rmtree(self.report_dir)
                logger.info(f"Cleared existing reports directory.")
            except Exception as e:
                logger.error(f"Error clearing reports directory: {e}")
        os.makedirs(self.report_dir, exist_ok=True)
        async with async_playwright() as p:
            browser, context, page = None, None, None
            try:
                browser = await p.chromium.launch(headless=False)
                context = await browser.new_context(
                    viewport={'width': 1460, 'height': 1080},
                    ignore_https_errors=True
                )
                await context.add_init_script("() => { Object.defineProperty(navigator, 'webdriver', { get: () => undefined }) }")
                page = await self.create_page(context)
                await page.goto(self.url, wait_until="networkidle", timeout=60000)
                await self.attempt_login(page)
                while urls_to_visit and pages_visited_count < max_pages:
                    current_url = urls_to_visit.pop(0)
                    if current_url.rstrip('/') in self.visited_urls:
                        continue
                    page, new_links = await self.crawl_url(context, page, current_url, base_domain)
                    pages_visited_count += 1
                    for link_info in new_links:
                        link_url = link_info['url'].rstrip('/')
                        if link_url not in self.visited_urls and link_url not in urls_to_visit:
                            if urlparse(link_url).netloc == base_domain:
                                urls_to_visit.append(link_url)
                                logger.debug(f"Added URL to visit: {link_url} (from {link_info.get('source','unknown')})")
                            else:
                                logger.debug(f"Skipping off-domain link: {link_url}")
                logger.info(f"Crawl loop finished. Pages visited: {pages_visited_count}")
            except Exception as e:
                logger.critical(f"Fatal error during Playwright crawl: {e}", exc_info=True)
            finally:
                if page and not page.is_closed():
                    await page.close()
                if context:
                    await context.close()
                if browser and browser.is_connected():
                    await browser.close()
        duration = datetime.now() - start_time
        logger.info(f"Crawl phase completed. Visited: {pages_visited_count}, Total processed: {len(self.section_details)}, Screenshots: {len(self.screenshot_paths)}, Duration: {duration}")
        crawl_result = {
            "base_url": self.url,
            "pages_visited_count": pages_visited_count,
            "screenshot_paths": self.screenshot_paths,
            "crawl_details": self.section_details,
            "crawl_duration": str(duration)
        }
        crawl_summary_path = os.path.join(self.debug_dir, f"{self.app_id}_crawl_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
        try:
            with open(crawl_summary_path, 'w', encoding='utf-8') as f:
                json.dump(crawl_result, f, indent=2, default=str)
            logger.info(f"Crawl summary saved: {crawl_summary_path}")
        except Exception as e:
            logger.error(f"Error saving crawl summary: {e}")
        return crawl_result

    async def crawl_and_process(self, max_pages=10):
        crawl_result = await self.crawl(max_pages)
        logger.info("Initializing OCR Image Processor...")
        self.image_processor = ImageToDocProcessor(
            image_folder=self.screenshot_dir,
            app_id=self.app_id
        )
        process_result = await self.image_processor.process()
        combined_result = {
            "crawl_summary": {
                k: crawl_result.get(k) for k in ["base_url", "pages_visited_count", "screenshot_paths", "crawl_duration"]
            },
            "image_processing_summary": process_result
        }
        combined_result["crawl_summary"]["screenshot_count"] = len(crawl_result.get("screenshot_paths", []))
        combined_summary_path = os.path.join(self.report_dir, f"{self.app_id}_combined_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
        try:
            with open(combined_summary_path, 'w', encoding='utf-8') as f:
                json.dump(combined_result, f, indent=2, default=str)
            logger.info(f"Combined summary saved: {combined_summary_path}")
        except Exception as e:
            logger.error(f"Error saving combined summary: {e}")
        return combined_result

async def main():
    config = {
        'url': os.getenv("POC_START_URL", ""),
        'section': 'login',
        'sections': [],
        'app_id': 'herokuapp_ocr_integrated',
        'username': os.getenv("POC_USERNAME", ""),
        'password': os.getenv("POC_PASSWORD", ""),
        'max_pages': 8
    }
    crawler = AgnoBrowserCrawler(
        url=config['url'],
        section=config['section'],
        sections=config['sections'],
        app_id=config['app_id'],
        username=config['username'],
        password=config['password']
    )
    try:
        result = await crawler.crawl_and_process(max_pages=config['max_pages'])
        print("\n=== Crawl (Playwright/Crawl4AI) and Process (OCR) Summary ===")
        if result:
            crawl_summary = result.get('crawl_summary', {})
            proc_summary = result.get('image_processing_summary', {})
            print(f"Base URL: {crawl_summary.get('base_url')}")
            print(f"Pages Visited: {crawl_summary.get('pages_visited_count', 'N/A')}")
            print(f"Screenshots Taken: {crawl_summary.get('screenshot_count', 'N/A')}")
            print(f"Crawl Duration: {crawl_summary.get('crawl_duration', 'N/A')}")
            print("-" * 20)
            print(f"Image Processor: {proc_summary.get('processor', 'N/A')}")
            print(f"Images Processed: {proc_summary.get('processed_image_count', 'N/A')}")
            if proc_summary.get('error'):
                print(f"Image Processing Error: {proc_summary.get('error')}")
            print(f"Reports Generated:")
            print(f"  - Markdown: {proc_summary.get('report_path') or 'Failed/Skipped'}")
            print(f"  - HTML: {proc_summary.get('html_report_path') or 'Failed/Skipped'}")
            print(f"  - DOCX: {proc_summary.get('docx_report_path') or 'Failed/Skipped'}")
        else:
            print("Crawl and process did not return a result.")
    except Exception as e:
        logger.critical(f"An error occurred in main execution: {e}", exc_info=True)
        print(f"\nAn error occurred during execution: {e}")

if __name__ == "__main__":
    asyncio.run(main())