"""Core Playwright crawler with structured page analysis and QA test generation.

The production path intentionally stays vision-free: DOM, accessibility metadata,
forms, sections, links, and interactive elements are analyzed directly with
Playwright. Ollama is optional and is used only for text summaries / test ideas.
"""

import asyncio
import json
import logging
import os
import random
import re
import uuid
from datetime import datetime
from functools import wraps
from urllib.parse import quote, urljoin, urlparse

import markdown2
import requests
from playwright.async_api import TimeoutError as PlaywrightTimeoutError, async_playwright
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


def retry_with_backoff(max_retries=2, base_delay=5, max_delay=60):
    def decorator(func):
        @wraps(func)
        async def wrapper(*args, **kwargs):
            for attempt in range(max_retries):
                try:
                    return await func(*args, **kwargs)
                except Exception as e:
                    if attempt == max_retries - 1:
                        logger.error(
                            f"{func.__name__} failed after {max_retries} attempts: {e}",
                            exc_info=True,
                        )
                        raise
                    delay = min(base_delay * (2**attempt) + random.uniform(0, 0.1), max_delay)
                    logger.warning(
                        f"{func.__name__} attempt {attempt + 1}/{max_retries} failed with {e}. "
                        f"Retrying in {delay:.2f}s..."
                    )
                    await asyncio.sleep(delay)

        return wrapper

    return decorator


class WebpageContent(BaseModel):
    url: str
    content: str
    summary: str
    timestamp: str
    section: str
    status: str
    interactive_elements: list = Field(default_factory=list)
    title: str = ""
    meta_description: str = ""
    headings: list = Field(default_factory=list)
    sections: list = Field(default_factory=list)
    forms: list = Field(default_factory=list)
    interaction_candidates: list = Field(default_factory=list)
    generated_test_cases: list = Field(default_factory=list)
    word_count: int = 0
    link_count: int = 0
    qa_risk_score: int = 0
    qa_risk_level: str = "Low"
    qa_risk_factors: list = Field(default_factory=list)
    console_errors: list = Field(default_factory=list)
    failed_requests: list = Field(default_factory=list)
    page_load_ms: int = 0
    accessibility_findings: list = Field(default_factory=list)
    api_requests: list = Field(default_factory=list)


class WebCrawler:
    """Crawl a site and produce structured, RAG-ready web/QA documentation."""

    def __init__(
        self,
        url,
        section="Home",
        sections=None,
        app_id="crawl",
        username="",
        password="",
        ollama_url="",
        ollama_model="",
        llm_base_url="",
        llm_api_key="",
        llm_model="portfolio-free",
        output_dir=".",
    ):
        self.url = url
        self.section = section
        self.sections = sections or []
        self.app_id = app_id
        self.username = username
        self.password = password
        self.ollama_url = ollama_url
        self.ollama_model = ollama_model
        self.llm_base_url = llm_base_url
        self.llm_api_key = llm_api_key
        self.llm_model = llm_model
        self.visited_urls = set()
        self.link_collection = {}
        self.content_collection = {}
        self.debug_dir = os.path.join(output_dir, "debug")
        self.report_dir = os.path.join(output_dir, "reports")
        for directory in [self.debug_dir, self.report_dir]:
            os.makedirs(directory, exist_ok=True)

    async def infer_section_from_url(self, url: str, page) -> str:
        logger.debug(f"Inferring section for URL: {url}")
        try:
            # Prefer standards-based navigation/breadcrumb semantics so the crawler
            # works across sites instead of depending on vendor-specific CSS classes.
            breadcrumb = await page.query_selector(
                '[aria-label*="breadcrumb" i], nav[aria-label*="breadcrumb" i], [role="navigation"] [aria-current="page"]'
            )
            if breadcrumb:
                section = (await breadcrumb.text_content() or "").strip()
                if section:
                    return section

            active_menu = await page.query_selector(
                'nav a[aria-current="page"], nav [aria-current="page"], '
                '[role="navigation"] a[aria-current="page"], '
                '[role="navigation"] [aria-current="page"]'
            )
            if active_menu:
                section = (await active_menu.text_content() or "").strip()
                if section:
                    return section

            title = await page.title()
            title_section = title.strip().split("|")[0].strip() if title else ""
        except Exception as e:
            logger.warning(f"Error extracting section from page content: {e}")

        path = urlparse(url).path.lower().rstrip('/')
        segments = [s for s in path.split('/') if s and not s.isdigit()]
        if segments:
            last_segment = re.sub(r'^(view|search|list|module)', '', segments[-1], flags=re.IGNORECASE)
            label = ' '.join(word.capitalize() for word in last_segment.split('-'))
            if label:
                return label
        if title_section:
            return title_section
        return 'Unknown'

    async def create_page(self, context):
        return await context.new_page()

    async def attempt_login(self, page):
        if not self.username or not self.password:
            return False
        logger.info(f"Attempting login for URL: {page.url}")
        username_sel = 'input[name="username"], input[id="username"], input[type="text"], input[placeholder*="username" i]'
        password_sel = 'input[name="password"], input[id="password"], input[type="password"], input[placeholder*="password" i]'
        submit_selector = 'button[type="submit"], input[type="submit"]'
        try:
            username_field = await page.wait_for_selector(username_sel, state="visible", timeout=10000)
            password_field = await page.wait_for_selector(password_sel, state="visible", timeout=10000)
            login_button = await page.query_selector(submit_selector)

            # Some applications use a plain button without type=submit. Prefer
            # buttons whose accessible/text label indicates authentication,
            # while avoiding any site-specific selector or route.
            if not login_button:
                buttons = await page.query_selector_all('button')
                login_terms = re.compile(r"\b(log[ -]?in|sign[ -]?in|authenticate|continue|submit)\b", re.I)
                for candidate in buttons:
                    label = (await candidate.text_content() or "").strip()
                    aria_label = (await candidate.get_attribute("aria-label") or "").strip()
                    if login_terms.search(label) or login_terms.search(aria_label):
                        login_button = candidate
                        break

            if username_field and password_field and login_button:
                initial_url = page.url
                await username_field.fill(self.username)
                await password_field.fill(self.password)
                await login_button.click()
                try:
                    try:
                        await page.wait_for_load_state("networkidle", timeout=20000)
                    except PlaywrightTimeoutError:
                        pass
                    # A generic login-success check: navigation occurred or the
                    # password field is no longer visible. No site-specific route
                    # or CSS class is assumed.
                    if page.url != initial_url:
                        return True
                    password_visible = await page.locator(password_sel).first.is_visible()
                    return not password_visible
                except PlaywrightTimeoutError:
                    try:
                        return not await page.locator(password_sel).first.is_visible()
                    except Exception:
                        return False
            return False
        except PlaywrightTimeoutError:
            logger.warning("Timeout finding login elements, skipping form login.")
            return False
        except Exception as e:
            logger.error(f"Error during login attempt: {e}", exc_info=True)
            return False

    async def close_popups(self, page):
        if page.is_closed():
            return
        popup_selectors = [
            '[aria-label*="close" i]', '[aria-label*="dismiss" i]', 'button:has-text("Accept")',
            'button:has-text(" Agree")', 'button:has-text("OK")', 'button:has-text("Got it")',
            'button.close', 'button.modal-close', '[class*="cookie"] button', '[id*="cookie"] button',
            'div[role="dialog"] button[class*="close"]',
        ]
        for selector in popup_selectors:
            try:
                elements = await page.query_selector_all(selector)
                for element in elements:
                    if await element.is_visible() and await element.is_enabled():
                        try:
                            await element.click(timeout=5000)
                            await page.wait_for_timeout(500)
                        except Exception as e:
                            logger.debug(f"Could not click popup '{selector}': {e}")
            except Exception as e:
                logger.debug(f"Error querying selector '{selector}': {e}")

    async def extract_page_links_playwright(self, page, base_url: str) -> list:
        links_found = []
        if page.is_closed():
            return links_found
        try:
            await page.wait_for_load_state('domcontentloaded', timeout=20000)
            await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
            await page.wait_for_timeout(500)
            link_elements = await page.query_selector_all(
                'a[href], button[onclick], [role="link"], [role="button"], [data-href], [data-url], '
                'nav a, nav button, menu a, menu button'
            )
            current_page_url_parsed = urlparse(page.url)
            for element in link_elements:
                try:
                    href = await element.get_attribute('href') or await element.get_attribute('data-href') or await element.get_attribute('data-url')
                    text = (await element.text_content() or "").strip()
                    if href and not href.startswith(('javascript:', '#', 'mailto:', 'tel:')):
                        absolute_url = urljoin(base_url, href)
                        parsed = urlparse(absolute_url)
                        if parsed.netloc == current_page_url_parsed.netloc:
                            normalized = parsed._replace(fragment="", query="").geturl().rstrip('/')
                            if normalized:
                                links_found.append({"url": normalized, "text": text or "Link", "source": "playwright"})
                except Exception as e:
                    logger.debug(f"Error parsing link: {e}")
            return list({link['url']: link for link in links_found}.values())
        except Exception as e:
            logger.error(f"Error extracting links: {e}", exc_info=True)
            return []

    async def extract_interactive_elements(self, page) -> list:
        """Extract controls without executing them; actions are intentionally non-destructive."""
        try:
            elements = await page.query_selector_all(
                'input, button, select, textarea, a[href], [role="button"], [role="link"], [onclick]'
            )
            result = []
            for element in elements:
                try:
                    tag = await element.evaluate('el => el.tagName.toLowerCase()')
                    element_type = await element.get_attribute('type') or tag
                    text = (await element.text_content() or '').strip()
                    result.append({
                        'tag': tag,
                        'type': element_type,
                        'id': await element.get_attribute('id') or '',
                        'name': await element.get_attribute('name') or '',
                        'text': text,
                        'aria-label': await element.get_attribute('aria-label') or '',
                        'placeholder': await element.get_attribute('placeholder') or '',
                        'required': (await element.get_attribute('required')) is not None,
                        'disabled': (await element.get_attribute('disabled')) is not None,
                        'role': await element.get_attribute('role') or '',
                    })
                except Exception as e:
                    logger.debug(f"Error extracting interactive element: {e}")
            return result
        except Exception as e:
            logger.error(f"Error extracting interactive elements: {e}", exc_info=True)
            return []

    async def extract_page_structure(self, page) -> dict:
        """Build structured DOM metadata inspired by the POC analyzers."""
        structure = {
            "title": "",
            "meta_description": "",
            "headings": [],
            "sections": [],
            "forms": [],
            "accessibility_findings": [],
        }
        try:
            structure["title"] = (await page.title() or "").strip()
        except Exception:
            pass
        try:
            meta = await page.query_selector('meta[name="description"]')
            if meta:
                structure["meta_description"] = (await meta.get_attribute("content") or "").strip()
        except Exception:
            pass

        try:
            headings = await page.query_selector_all('h1, h2, h3, h4, h5, h6')
            for heading in headings:
                text = re.sub(r'\s+', ' ', (await heading.text_content() or '').strip())
                if text:
                    structure["headings"].append({
                        "level": await heading.evaluate('el => parseInt(el.tagName.substring(1), 10)'),
                        "text": text,
                    })
        except Exception as e:
            logger.debug(f"Heading extraction failed: {e}")

        try:
            section_nodes = await page.query_selector_all('main > section, article > section, section, article')
            seen = set()
            for node in section_nodes:
                text = re.sub(r'\s+', ' ', (await node.text_content() or '').strip())
                if not text or text in seen:
                    continue
                seen.add(text)
                heading = await node.query_selector('h1, h2, h3, h4, h5, h6')
                heading_text = re.sub(r'\s+', ' ', (await heading.text_content() or '').strip()) if heading else ''
                structure["sections"].append({
                    "heading": heading_text,
                    "text": text[:1000],
                    "word_count": len(text.split()),
                })
                if len(structure["sections"]) >= 30:
                    break
        except Exception as e:
            logger.debug(f"Section extraction failed: {e}")

        try:
            forms = await page.query_selector_all('form')
            for form_index, form in enumerate(forms, 1):
                fields = []
                field_nodes = await form.query_selector_all('input, select, textarea')
                for field in field_nodes:
                    field_type = await field.get_attribute('type') or await field.evaluate('el => el.tagName.toLowerCase()')
                    fields.append({
                        "tag": await field.evaluate('el => el.tagName.toLowerCase()'),
                        "type": field_type,
                        "name": await field.get_attribute('name') or '',
                        "id": await field.get_attribute('id') or '',
                        "placeholder": await field.get_attribute('placeholder') or '',
                        "required": (await field.get_attribute('required')) is not None,
                    })
                action = await form.get_attribute('action') or ''
                method = (await form.get_attribute('method') or 'get').lower()
                submit_nodes = await form.query_selector_all('button[type="submit"], input[type="submit"], button:not([type])')
                structure["forms"].append({
                    "index": form_index,
                    "action": action,
                    "method": method,
                    "field_count": len(fields),
                    "fields": fields,
                    "submit_control_count": len(submit_nodes),
                })
        except Exception as e:
            logger.debug(f"Form extraction failed: {e}")
        return structure

    async def assess_accessibility(self, page) -> list:
        """Run lightweight, DOM-based accessibility checks without claiming WCAG certification."""
        findings = []
        try:
            title = (await page.title() or '').strip()
            if not title:
                findings.append({"severity": "Medium", "rule": "page-title", "message": "Page has no document title."})
            lang = await page.get_attribute("html", "lang")
            if not lang:
                findings.append({"severity": "Low", "rule": "html-lang", "message": "Document language is not declared."})
            images = await page.query_selector_all("img")
            missing_alt = 0
            for image in images:
                if (await image.get_attribute("alt")) is None:
                    missing_alt += 1
            if missing_alt:
                findings.append({"severity": "Medium", "rule": "image-alt", "message": f"{missing_alt} image(s) are missing alt attributes."})
            controls = await page.query_selector_all("input, select, textarea, button")
            unlabeled = 0
            for control in controls:
                if await control.get_attribute("type") == "hidden":
                    continue
                aria = (await control.get_attribute("aria-label") or '').strip()
                labelledby = (await control.get_attribute("aria-labelledby") or '').strip()
                text = (await control.text_content() or '').strip()
                cid = (await control.get_attribute("id") or '').strip()
                has_label = bool(aria or labelledby or text)
                if not has_label and cid:
                    label = await page.query_selector(f'label[for="{cid}"]')
                    has_label = label is not None
                if not has_label:
                    unlabeled += 1
            if unlabeled:
                findings.append({"severity": "High", "rule": "form-label", "message": f"{unlabeled} interactive control(s) have no detectable accessible name/label."})
            headings = await page.query_selector_all("h1, h2, h3, h4, h5, h6")
            levels = [int(await h.evaluate("el => el.tagName.substring(1)")) for h in headings]
            if levels and levels[0] != 1:
                findings.append({"severity": "Low", "rule": "heading-order", "message": "The first heading is not an H1."})
            if not headings:
                findings.append({"severity": "Medium", "rule": "heading-structure", "message": "No semantic headings were detected."})
            return findings[:50]
        except Exception as exc:
            logger.debug("Accessibility assessment failed: %s", exc)
            return findings

    def build_interaction_candidates(self, interactive_elements: list, forms: list) -> list:
        candidates = []
        for element in interactive_elements:
            tag = element.get('tag', '')
            element_type = element.get('type', '')
            label = element.get('aria-label') or element.get('text') or element.get('name') or element.get('id') or element.get('placeholder') or tag
            if element.get('disabled'):
                continue
            if tag in ('input', 'textarea', 'select'):
                action = 'fill/select value'
            elif tag == 'a' or element.get('role') == 'link':
                action = 'navigate'
            else:
                action = 'click'
            candidates.append({
                "action": action,
                "target": label[:120],
                "type": element_type,
                "safe_by_default": action == 'navigate',
            })
        for form in forms:
            candidates.append({
                "action": "submit form",
                "target": form.get('action') or f"form #{form.get('index')}",
                "type": form.get('method', 'get').upper(),
                "safe_by_default": False,
            })
        return candidates

    def assess_qa_risk(self, page_data: dict) -> dict:
        """Score QA risk from observable page characteristics, not guessed business logic."""
        forms = page_data.get("forms", [])
        elements = page_data.get("interactive_elements", [])
        headings = page_data.get("headings", [])
        console_errors = page_data.get("console_errors", [])
        failed_requests = page_data.get("failed_requests", [])
        accessibility_findings = page_data.get("accessibility_findings", [])
        score = 5
        factors = []

        if forms:
            score += min(25, 10 + len(forms) * 5)
            factors.append(f"{len(forms)} form(s) detected")
        required_fields = sum(1 for form in forms for field in form.get("fields", []) if field.get("required"))
        if required_fields:
            score += min(15, required_fields * 3)
            factors.append(f"{required_fields} required field(s)")
        password_fields = sum(1 for e in elements if str(e.get("type", "")).lower() == "password")
        if password_fields:
            score += 15
            factors.append("authentication-sensitive input detected")
        buttons = sum(1 for e in elements if e.get("tag") == "button" or e.get("role") == "button")
        links = sum(1 for e in elements if e.get("tag") == "a" or e.get("role") == "link")
        if buttons >= 5:
            score += min(10, buttons)
            factors.append(f"{buttons} interactive button(s)")
        if links >= 25:
            score += 5
            factors.append(f"{links} navigation link(s)")
        if not headings:
            score += 5
            factors.append("no heading structure detected")
        if console_errors:
            score += min(15, len(console_errors) * 3)
            factors.append(f"{len(console_errors)} browser console error(s)")
        if failed_requests:
            score += min(15, len(failed_requests) * 3)
            factors.append(f"{len(failed_requests)} failed network request(s)")
        high_a11y = sum(1 for finding in accessibility_findings if finding.get("severity") == "High")
        if accessibility_findings:
            score += min(10, len(accessibility_findings) * 2)
            factors.append(f"{len(accessibility_findings)} accessibility finding(s)")
        if high_a11y:
            score += min(8, high_a11y * 4)

        score = min(100, score)
        level = "High" if score >= 70 else "Medium" if score >= 40 else "Low"
        return {"score": score, "level": level, "factors": factors}

    def generate_test_cases(self, page_data: dict) -> list:
        """Generate evidence-grounded QA cases from discovered DOM capabilities."""
        cases = []
        url = page_data.get('url', '')
        forms = page_data.get('forms', [])
        elements = page_data.get('interactive_elements', [])
        headings = page_data.get('headings', [])
        console_errors = page_data.get('console_errors', [])
        failed_requests = page_data.get('failed_requests', [])
        accessibility_findings = page_data.get('accessibility_findings', [])

        def add(title, objective, priority='Medium', category='Functional', steps=None, expected=None, evidence=None, execution='manual_or_safe_automation'):
            cases.append({
                "id": f"TC-{len(cases) + 1:03d}",
                "title": title,
                "objective": objective,
                "priority": priority,
                "category": category,
                "url": url,
                "preconditions": ["Page is reachable and loaded successfully."],
                "steps": steps or ["Open the page.", "Perform the described verification."],
                "expected_result": expected or "The observed behavior matches the expected application behavior without errors.",
                "evidence": evidence or [f"Observed on {url}"],
                "execution": execution,
            })

        if forms:
            add("Verify form renders correctly", "Verify all detected form fields and submit controls are visible and usable.", "High", "Functional", ["Open the page.", "Locate each detected form.", "Verify fields and submit controls are visible and enabled."], "All detected controls are rendered, labeled, and usable.", [f"Detected {len(forms)} form(s)"])
            required_fields = sum(1 for form in forms for field in form.get('fields', []) if field.get('required'))
            if required_fields:
                add("Verify required-field validation", f"Verify required-field validation for {required_fields} required field(s).", "High", "Validation", ["Open the form.", "Leave required fields empty.", "Attempt the form action without entering those values."], "Clear validation feedback is shown and invalid submission is prevented or handled correctly.", [f"Detected {required_fields} required field(s)"])
            add("Verify form submission", "Verify a form can be submitted with representative valid test data and reaches the expected application state.", "High", "Functional", ["Open the form.", "Enter safe test data appropriate to each field type.", "Submit using an authorized test environment.", "Verify the resulting page or response."], "The form completes successfully or presents actionable validation feedback.", [f"Detected {len(forms)} form(s) with {sum(f.get('field_count', 0) for f in forms)} field(s)"], "manual_or_authorized_automation")

        input_count = sum(1 for e in elements if e.get('tag') in ('input', 'textarea', 'select'))
        if input_count:
            add("Verify input controls", f"Verify {input_count} detected input/select/textarea control(s) accept appropriate values and expose usable labels.", "Medium", "Accessibility", ["Inspect each input control.", "Verify its label, placeholder, type, and state.", "Enter a representative value where safe."], "Controls accept appropriate values and provide an understandable accessible name or label.", [f"Detected {input_count} input control(s)"])

        click_count = sum(1 for e in elements if e.get('tag') == 'button' or e.get('role') == 'button')
        if click_count:
            add("Verify interactive buttons", f"Verify {click_count} detected button control(s) respond correctly to authorized user interaction.", "Medium", "Functional", ["Identify each button.", "Verify its label and enabled state.", "Activate it in a safe test environment.", "Verify the resulting state."], "The button performs its intended action or provides clear feedback.", [f"Detected {click_count} button control(s)"], "manual_or_authorized_automation")

        link_count = sum(1 for e in elements if e.get('tag') == 'a' or e.get('role') == 'link')
        if link_count:
            add("Verify navigation links", f"Verify {link_count} detected navigation link(s) resolve to intended destinations.", "Medium", "Navigation", ["Open each internal link.", "Verify the destination loads successfully.", "Check for unexpected redirects or errors."], "Links resolve to reachable, expected destinations.", [f"Detected {link_count} navigation link(s)"])

        if headings:
            add("Verify page structure", "Verify the page exposes a meaningful heading hierarchy and section structure.", "Low", "Content", ["Inspect the heading hierarchy.", "Check for a logical H1-to-Hn structure.", "Verify major content sections have meaningful labels."], "The page structure is understandable and logically organized.", [f"Detected {len(headings)} heading(s)"])
        else:
            add("Review missing heading structure", "Review whether the page needs semantic headings for usability and accessibility.", "Medium", "Accessibility", ["Inspect the main content area.", "Determine whether major sections have semantic headings."], "Important page sections have meaningful semantic headings where appropriate.", ["No H1-H6 headings detected"])

        if accessibility_findings:
            add("Review accessibility findings", f"Review {len(accessibility_findings)} DOM-based accessibility finding(s) and confirm remediation where appropriate.", "High" if any(f.get("severity") == "High" for f in accessibility_findings) else "Medium", "Accessibility", ["Open the page.", "Review the flagged controls or document metadata.", "Verify each issue against the application's accessibility requirements."], "Interactive controls and document semantics expose appropriate accessible names and structure.", [f"Captured {len(accessibility_findings)} accessibility finding(s)"], "manual_or_safe_automation")

        if console_errors:
            add("Investigate browser console errors", f"Investigate {len(console_errors)} browser console error(s) captured during page load.", "High", "Reliability", ["Open the page.", "Reproduce the load state.", "Inspect console errors and their source.", "Determine whether they affect user-visible behavior."], "No unexpected application errors remain in the browser console.", [f"Captured {len(console_errors)} console error(s)"])
        if failed_requests:
            add("Investigate failed network requests", f"Investigate {len(failed_requests)} failed network request(s) captured during page load.", "High", "Reliability", ["Open the page.", "Inspect failed requests.", "Verify response status and request URL.", "Determine whether each failure is expected or a defect."], "Critical resources and application requests complete successfully.", [f"Captured {len(failed_requests)} failed request(s)"])
        return cases

    async def summarize_content(self, content: str) -> str:
        if not content:
            return ""
        if self.llm_base_url and self.llm_api_key:
            return await self._summarize_via_gateway(content)
        if self.ollama_url:
            return await self._summarize_via_ollama(content)
        return content[:200] + "..." if len(content) > 200 else content

    @retry_with_backoff(max_retries=2, base_delay=2, max_delay=15)
    async def _summarize_via_gateway(self, content: str) -> str:
        """Summarize through the portfolio's OpenAI-compatible LiteLLM gateway."""
        try:
            prompt = (
                "You are an expert web application analyst. Summarize the page in 50-150 words. "
                "Focus on its purpose, user actions, forms, important data, and functionality. "
                "Do not discuss visual appearance or screenshots. Content:\n\n" + content
            )
            prompt = prompt[:4000].rsplit(" ", 1)[0] + "..." if len(prompt) > 4000 else prompt
            payload = {
                "model": self.llm_model,
                "messages": [
                    {"role": "system", "content": "Summarize only the supplied webpage content."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.3,
                "max_tokens": 220,
            }
            response = requests.post(
                f"{self.llm_base_url.rstrip('/')}/chat/completions",
                headers={"Authorization": f"Bearer {self.llm_api_key}"},
                json=payload,
                timeout=30,
            )
            response.raise_for_status()
            result = response.json()
            summary = result.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
            if not summary:
                return content[:200] + "..." if content else ""
            return summary
        except Exception as e:
            logger.error(f"Error summarizing content via LLM gateway ({self.llm_model}): {e}")
            return content[:200] + "..." if content else ""

    @retry_with_backoff(max_retries=2, base_delay=5, max_delay=60)
    async def _summarize_via_ollama(self, content: str) -> str:
        try:
            prompt = (
                "You are an expert web application analyst. Summarize the page in 50-150 words. "
                "Focus on its purpose, user actions, forms, important data, and functionality. "
                "Do not discuss visual appearance or screenshots. Content:\n\n" + content
            )
            max_input_chars = 4000
            if len(prompt) > max_input_chars:
                prompt = prompt[:max_input_chars].rsplit(' ', 1)[0] + '...'
            payload = {
                "model": self.ollama_model,
                "prompt": prompt,
                "stream": False,
                "options": {"max_tokens": 200, "temperature": 0.7},
            }
            response = requests.post(self.ollama_url, json=payload, timeout=30)
            response.raise_for_status()
            result = response.json()
            summary = result.get("response", "").strip()
            if not summary:
                return content[:200] + "..." if content else ""
            debug_path = os.path.join(self.debug_dir, f"ollama_summary_{uuid.uuid4().hex[:8]}.txt")
            with open(debug_path, 'w', encoding='utf-8') as f:
                f.write(f"---PROMPT---\n{prompt}\n\n---SUMMARY---\n{summary}")
            return summary
        except Exception as e:
            logger.error(f"Error summarizing content with Ollama ({self.ollama_model}): {e}")
            return content[:200] + "..." if content else ""

    @retry_with_backoff(max_retries=2, base_delay=5, max_delay=60)
    async def scrape_page_content(self, url: str, section: str, page) -> WebpageContent:
        logger.info(f"Scraping content for: {url}")
        content = ""
        summary = ""
        status = "failed: no content"
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        interactive_elements = []
        structure = {"title": "", "meta_description": "", "headings": [], "sections": [], "forms": []}
        console_errors = []
        failed_requests = []
        api_requests = []
        load_started = asyncio.get_running_loop().time()
        try:
            def on_console(msg):
                if msg.type == "error":
                    text = (msg.text or "").strip()
                    if text and text not in console_errors and len(console_errors) < 50:
                        console_errors.append(text)

            def on_request_failed(request):
                if len(failed_requests) < 50:
                    failed_requests.append({"url": request.url, "method": request.method, "failure": request.failure})

            def on_response(response):
                try:
                    resource = response.request.resource_type
                    content_type = (response.headers.get("content-type") or "").lower()
                    if resource in {"xhr", "fetch"} or "json" in content_type:
                        if len(api_requests) < 100:
                            api_requests.append({
                                "method": response.request.method,
                                "url": response.url,
                                "status": response.status,
                                "resource_type": resource,
                            })
                except Exception:
                    pass

            page.on("console", on_console)
            page.on("requestfailed", on_request_failed)
            page.on("response", on_response)
            await page.goto(url, wait_until="domcontentloaded", timeout=30000)
            # Do not wait for networkidle: modern apps can keep background sockets,
            # polling, analytics, or long-lived requests open indefinitely. Playwright
            # recommends relying on explicit readiness/DOM assertions instead.
            await page.wait_for_timeout(500)
            # If credentials were supplied and this page still exposes a
            # password field, attempt the generic login flow for this page.
            if self.username and self.password:
                password_locator = page.locator('input[type="password"], input[name*="password" i], input[id*="password" i]')
                if await password_locator.count():
                    await self.attempt_login(page)

            structure = await self.extract_page_structure(page)
            structure["accessibility_findings"] = await self.assess_accessibility(page)
            content_element = await page.query_selector(
                'main, article, [role="main"], div#content, div.container'
            )
            if content_element:
                content = (await content_element.text_content() or "").strip()
                content = re.sub(r"\s+", " ", content).strip()
            if not content:
                content = await page.evaluate("""
                    () => {
                        const exclude = ['nav','footer','style','script','[role="alert"]','[class*="cookie"]'];
                        const selector = exclude.join(',');
                        const clone = document.body.cloneNode(true);
                        clone.querySelectorAll(selector).forEach(el => el.remove());
                        return (clone.innerText || '').replace(/\\s+/g, ' ').trim();
                    }
                """)
            if content:
                status = "success: Playwright"
            interactive_elements = await self.extract_interactive_elements(page)
        except Exception as e:
            logger.error(f"Scrape failed for {url}: {e}", exc_info=True)
            status = f"failed: {str(e)}"

        interaction_candidates = self.build_interaction_candidates(interactive_elements, structure.get('forms', []))
        page_load_ms = int((asyncio.get_running_loop().time() - load_started) * 1000)
        page_data = {
            "url": url,
            "interactive_elements": interactive_elements,
            "forms": structure.get('forms', []),
            "headings": structure.get('headings', []),
            "console_errors": console_errors,
            "failed_requests": failed_requests,
            "accessibility_findings": structure.get("accessibility_findings", []),
        }
        risk = self.assess_qa_risk(page_data)
        test_cases = self.generate_test_cases(page_data)
        if content:
            summary = await self.summarize_content(content)

        return WebpageContent(
            url=url,
            content=content,
            summary=summary,
            timestamp=timestamp,
            section=section,
            status=status,
            interactive_elements=interactive_elements,
            title=structure.get('title', ''),
            meta_description=structure.get('meta_description', ''),
            headings=structure.get('headings', []),
            sections=structure.get('sections', []),
            forms=structure.get('forms', []),
            interaction_candidates=interaction_candidates,
            generated_test_cases=test_cases,
            word_count=len(content.split()),
            link_count=len([e for e in interactive_elements if e.get('tag') == 'a']),
            qa_risk_score=risk["score"],
            qa_risk_level=risk["level"],
            qa_risk_factors=risk["factors"],
            console_errors=console_errors,
            failed_requests=failed_requests,
            page_load_ms=page_load_ms,
            accessibility_findings=structure.get("accessibility_findings", []),
            api_requests=api_requests,
        )

    async def crawl_url(self, context, page, url_to_crawl: str, base_domain: str, section: str):
        normalized_url = url_to_crawl.rstrip('/')
        if normalized_url in self.visited_urls:
            return page, []
        logger.info(f"Crawling URL: {url_to_crawl}")
        combined_links = []
        try:
            target_goto_url = url_to_crawl
            parsed = urlparse(url_to_crawl)
            if "/basic_auth" in parsed.path.lower() and "@" not in parsed.netloc and self.username and self.password:
                target_goto_url = (
                    f"{parsed.scheme}://{quote(self.username, safe='')}:{quote(self.password, safe='')}@"
                    f"{parsed.netloc}{parsed.path}"
                )
                if parsed.query:
                    target_goto_url += f"?{parsed.query}"
            await page.goto(target_goto_url, wait_until="domcontentloaded", timeout=30000)
            await page.wait_for_timeout(500)
            final_url = page.url.rstrip('/')
            if final_url != normalized_url and urlparse(final_url).path != urlparse(normalized_url).path:
                if final_url in self.visited_urls:
                    return page, []
                normalized_url = final_url
            self.visited_urls.add(normalized_url)
            await self.close_popups(page)
            inferred_section = await self.infer_section_from_url(normalized_url, page)
            page_base_url = urlparse(page.url)._replace(path="", query="", fragment="").geturl()
            combined_links = await self.extract_page_links_playwright(page, page_base_url)
            for link in combined_links:
                self.link_collection[link['url']] = {
                    "text": link['text'], "source": link['source'], "section": inferred_section,
                }
            content = await self.scrape_page_content(normalized_url, inferred_section, page)
            self.content_collection[normalized_url] = content
            self.link_collection[normalized_url] = {
                "text": inferred_section or "Homepage", "source": "initial", "section": inferred_section,
            }
        except Exception as e:
            logger.error(f"Error crawling URL {url_to_crawl}: {e}", exc_info=True)
            self.visited_urls.add(normalized_url)
            self.content_collection[normalized_url] = WebpageContent(
                url=normalized_url, content="", summary="", timestamp=datetime.now().strftime("%Y%m%d_%H%M%S"),
                section=section, status=f"failed: {str(e)}", interactive_elements=[],
            )
        return page, combined_links

    async def crawl(self, max_pages=20, headless=True, progress_cb=None):
        logger.info(f"Starting crawl for base URL: {self.url}, Max Pages: {max_pages}")
        start_time = datetime.now()
        urls_to_visit = []
        initial_urls = [self.url] + [urljoin(self.url, s) for s in self.sections]
        for u in initial_urls:
            normalized = u.rstrip('/')
            if normalized not in self.visited_urls:
                urls_to_visit.append((normalized, self.section))
        pages_visited_count = 0
        base_domain = urlparse(self.url).netloc
        fatal_error = None
        async with async_playwright() as p:
            browser, context, page = None, None, None
            try:
                browser = await p.chromium.launch(headless=headless)
                context = await browser.new_context(viewport={'width': 1460, 'height': 1080}, ignore_https_errors=True)
                await context.add_init_script("() => { Object.defineProperty(navigator, 'webdriver', { get: () => undefined }) }")
                page = await self.create_page(context)
                await page.goto(self.url, wait_until="domcontentloaded", timeout=30000)
                await page.wait_for_timeout(500)
                login_success = await self.attempt_login(page)
                # Do not redirect to a site-specific dashboard. Continue from
                # the supplied start URL after authentication succeeds.
                while urls_to_visit and pages_visited_count < max_pages:
                    current_url, section = urls_to_visit.pop(0)
                    if current_url.rstrip('/') in self.visited_urls:
                        continue
                    page, new_links = await self.crawl_url(context, page, current_url, base_domain, section)
                    pages_visited_count += 1
                    if progress_cb:
                        progress_cb(pages_visited_count, max_pages, current_url)
                    queued = {u[0] for u in urls_to_visit}
                    for link_info in new_links:
                        link_url = link_info['url'].rstrip('/')
                        if link_url not in self.visited_urls and link_url not in queued and urlparse(link_url).netloc == base_domain:
                            urls_to_visit.append((link_url, link_info.get('section') or section))
            except Exception as e:
                logger.critical(f"Fatal error during Playwright crawl: {e}", exc_info=True)
                fatal_error = e
            finally:
                if page and not page.is_closed():
                    await page.close()
                if context:
                    await context.close()
                if browser and browser.is_connected():
                    await browser.close()
        if fatal_error is not None:
            raise fatal_error

        duration = datetime.now() - start_time
        pages = list(self.content_collection.values())
        crawl_result = {
            "base_url": self.url,
            "pages_visited_count": pages_visited_count,
            "link_count": len(self.link_collection),
            "content_count": len(self.content_collection),
            "successful_scrapes": len([c for c in pages if c.status.startswith("success")]),
            "failed_scrapes": len([c for c in pages if not c.status.startswith("success")]),
            "form_count": sum(len(c.forms) for c in pages),
            "interactive_element_count": sum(len(c.interactive_elements) for c in pages),
            "test_case_count": sum(len(c.generated_test_cases) for c in pages),
            "section_count": sum(len(c.sections) for c in pages),
            "high_risk_pages": sum(c.qa_risk_level == "High" for c in pages),
            "medium_risk_pages": sum(c.qa_risk_level == "Medium" for c in pages),
            "low_risk_pages": sum(c.qa_risk_level == "Low" for c in pages),
            "console_error_count": sum(len(c.console_errors) for c in pages),
            "failed_request_count": sum(len(c.failed_requests) for c in pages),
            "accessibility_finding_count": sum(len(c.accessibility_findings) for c in pages),
            "api_request_count": sum(len(c.api_requests) for c in pages),
            "crawl_duration": str(duration),
        }
        crawl_summary_path = os.path.join(self.debug_dir, f"{self.app_id}_crawl_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
        with open(crawl_summary_path, 'w', encoding='utf-8') as f:
            json.dump(crawl_result, f, indent=2, default=str)
        return crawl_result

    def _markdown_test_cases(self, test_cases: list) -> str:
        if not test_cases:
            return "No test cases generated.\n"
        lines = ["| ID | Test case | Priority | Category |", "|---|---|---|---|"]
        for case in test_cases:
            lines.append(f"| {case['id']} | {case['title']} | {case['priority']} | {case['category']} |")
        lines.append("")
        for case in test_cases:
            lines.extend([
                f"### {case['id']} — {case['title']}",
                f"**Objective:** {case['objective']}",
                f"**Priority:** {case['priority']}  ",
                f"**Category:** {case['category']}",
                "**Preconditions:**",
                *[f"- {item}" for item in case.get('preconditions', [])],
                "**Steps:**",
                *[f"{i}. {step}" for i, step in enumerate(case.get('steps', []), 1)],
                f"**Expected result:** {case.get('expected_result', '')}",
                "**Evidence:**",
                *[f"- {item}" for item in case.get('evidence', [])],
                "",
            ])
        return "\n".join(lines) + "\n"

    def _build_regression_snapshot(self) -> dict:
        """Create a stable, non-content snapshot for lightweight regression comparison."""
        pages = {}
        for page in self.content_collection.values():
            pages[page.url] = {
                "title": page.title,
                "status": page.status.split(":", 1)[0],
                "form_count": len(page.forms),
                "interactive_element_count": len(page.interactive_elements),
                "heading_count": len(page.headings),
                "link_count": page.link_count,
                "qa_risk_level": page.qa_risk_level,
                "accessibility_finding_count": len(page.accessibility_findings),
                "api_request_count": len(page.api_requests),
            }
        return {"base_url": self.url, "pages": pages}

    def compare_with_baseline(self) -> dict:
        """Compare this crawl with the previous snapshot, if one exists."""
        baseline_path = os.path.join(self.report_dir, f"{self.app_id}_baseline.json")
        current = self._build_regression_snapshot()
        if not os.path.exists(baseline_path):
            with open(baseline_path, "w", encoding="utf-8") as f:
                json.dump(current, f, indent=2)
            return {"available": False, "message": "Baseline created for future regression comparisons.", "added": [], "removed": [], "changed": []}
        try:
            with open(baseline_path, "r", encoding="utf-8") as f:
                previous = json.load(f)
        except (OSError, json.JSONDecodeError):
            previous = {"pages": {}}
        old_pages = previous.get("pages", {})
        new_pages = current.get("pages", {})
        added = sorted(set(new_pages) - set(old_pages))
        removed = sorted(set(old_pages) - set(new_pages))
        changed = []
        for url in sorted(set(old_pages) & set(new_pages)):
            if old_pages[url] != new_pages[url]:
                changed.append({"url": url, "before": old_pages[url], "after": new_pages[url]})
        # Update baseline only after comparison succeeds.
        with open(baseline_path, "w", encoding="utf-8") as f:
            json.dump(current, f, indent=2)
        return {
            "available": True,
            "message": "Compared with previous crawl baseline.",
            "added": added,
            "removed": removed,
            "changed": changed,
        }

    def generate_rag_document(self):
        report_base_name = f"{self.app_id}_rag_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        all_cases = [case for page in self.content_collection.values() for case in page.generated_test_cases]
        md_content = "# Web Crawl & QA Analysis\n\n"
        md_content += f"**Generated on**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        md_content += f"**Application ID**: {self.app_id}\n"
        md_content += f"**Base URL**: {self.url}\n\n"
        if not self.content_collection:
            md_content += "## No Content\n\nNo webpages were successfully scraped.\n"
        else:
            pages = list(self.content_collection.values())
            md_content += "## Crawl Summary\n\n"
            md_content += f"- **Pages visited**: {len(pages)}\n"
            md_content += f"- **Successful**: {sum(c.status.startswith('success') for c in pages)}\n"
            md_content += f"- **Forms**: {sum(len(c.forms) for c in pages)}\n"
            md_content += f"- **Interactive elements**: {sum(len(c.interactive_elements) for c in pages)}\n"
            risk_counts = {"High": sum(c.qa_risk_level == "High" for c in pages), "Medium": sum(c.qa_risk_level == "Medium" for c in pages), "Low": sum(c.qa_risk_level == "Low" for c in pages)}
            md_content += f"- **Generated test cases**: {len(all_cases)}\n"
            md_content += f"- **QA risk**: {risk_counts['High']} high / {risk_counts['Medium']} medium / {risk_counts['Low']} low\n"
            md_content += f"- **Console errors**: {sum(len(c.console_errors) for c in pages)}\n"
            md_content += f"- **Failed network requests**: {sum(len(c.failed_requests) for c in pages)}\n"
            md_content += f"- **Accessibility findings**: {sum(len(c.accessibility_findings) for c in pages)}\n"
            md_content += f"- **API/XHR responses**: {sum(len(c.api_requests) for c in pages)}\n\n"

            md_content += "## Regression Snapshot\n\n"
            md_content += "The crawl stores a lightweight baseline of page structure, forms, interactions, links, and QA risk for future comparisons.\n\n"
            md_content += "## Generated QA Test Plan\n\n"
            md_content += self._markdown_test_cases(all_cases) + "\n"

            for content_obj in sorted(pages, key=lambda c: c.timestamp):
                md_content += f"## Page: {content_obj.section or 'Homepage'}\n\n"
                md_content += f"**URL**: {content_obj.url}\n\n"
                md_content += f"**Title**: {content_obj.title or 'N/A'}\n\n"
                if content_obj.meta_description:
                    md_content += f"**Meta description**: {content_obj.meta_description}\n\n"
                md_content += f"**Status**: {content_obj.status}\n\n"
                md_content += f"**QA risk**: {content_obj.qa_risk_level} ({content_obj.qa_risk_score}/100)\n\n"
                if content_obj.qa_risk_factors:
                    md_content += "**Risk factors:**\n" + "\n".join(f"- {x}" for x in content_obj.qa_risk_factors) + "\n\n"
                md_content += f"**Page load**: {content_obj.page_load_ms} ms\n\n"
                md_content += f"**Browser console errors**: {len(content_obj.console_errors)}\n\n"
                md_content += f"**Failed network requests**: {len(content_obj.failed_requests)}\n\n"
                md_content += f"**Accessibility findings**: {len(content_obj.accessibility_findings)}\n\n"
                md_content += f"**API/XHR responses**: {len(content_obj.api_requests)}\n\n"
                md_content += f"**Summary**:\n{content_obj.summary or 'No summary generated'}\n\n"
                if content_obj.headings:
                    md_content += "**Headings**:\n"
                    for heading in content_obj.headings:
                        md_content += f"- H{heading['level']}: {heading['text']}\n"
                    md_content += "\n"
                if content_obj.forms:
                    md_content += "**Forms**:\n\n"
                    for form in content_obj.forms:
                        md_content += f"- Form {form['index']}: {form['method'].upper()} {form['action'] or '(current page)'} — {form['field_count']} fields\n"
                        for field in form['fields']:
                            required = "required" if field['required'] else "optional"
                            md_content += f"  - {field['type']} `{field['name'] or field['id'] or field['placeholder'] or 'unnamed'}` ({required})\n"
                    md_content += "\n"
                md_content += "**Interaction candidates**:\n"
                for candidate in content_obj.interaction_candidates[:100]:
                    md_content += f"- {candidate['action']}: {candidate['target']} (safe-by-default: {candidate['safe_by_default']})\n"
                md_content += "\n**Generated test cases**:\n\n"
                md_content += self._markdown_test_cases(content_obj.generated_test_cases) + "\n"
                md_content += f"**Content**:\n{content_obj.content or 'No content extracted'}\n\n---\n\n"

        md_path = os.path.join(self.report_dir, f"{report_base_name}.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        html_path = os.path.join(self.report_dir, f"{report_base_name}.html")
        html_body = markdown2.markdown(md_content, extras=["tables", "fenced-code-blocks"])
        html_full = (
            '<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8">'
            f'<title>Web Crawl QA Report - {self.app_id}</title>'
            '<style>body{font-family:sans-serif;margin:20px;line-height:1.6} '
            'table{border-collapse:collapse;width:100%} th,td{border:1px solid #ddd;padding:6px} '
            'pre{background:#f0f0f0;padding:10px}</style></head>'
            f'<body>{html_body}</body></html>'
        )
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_full)
        return md_path, html_path

    def generate_docx_report(self):
        """Generate a portable DOCX report using the structured crawl data."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx is not installed; skipping DOCX export")
            return None

        report_base_name = f"{self.app_id}_qa_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        path = os.path.join(self.report_dir, f"{report_base_name}.docx")
        pages = list(self.content_collection.values())
        all_cases = [case for page in pages for case in page.generated_test_cases]
        doc = Document()
        doc.add_heading("Web Crawl & QA Analysis", 0)
        doc.add_paragraph(f"Base URL: {self.url}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        risk_counts = {"High": sum(p.qa_risk_level == "High" for p in pages), "Medium": sum(p.qa_risk_level == "Medium" for p in pages), "Low": sum(p.qa_risk_level == "Low" for p in pages)}
        doc.add_heading("QA Overview", level=1)
        doc.add_paragraph(f"Risk distribution: {risk_counts['High']} high, {risk_counts['Medium']} medium, {risk_counts['Low']} low.")
        doc.add_paragraph(f"Browser console errors: {sum(len(p.console_errors) for p in pages)} | Failed requests: {sum(len(p.failed_requests) for p in pages)}")
        doc.add_heading("QA Test Plan", level=1)
        for case in all_cases:
            doc.add_heading(f"{case['id']} — {case['title']}", level=2)
            doc.add_paragraph(f"{case['objective']} | Priority: {case['priority']} | Category: {case['category']}")
            doc.add_paragraph("Steps:")
            for step in case.get("steps", []):
                doc.add_paragraph(step, style="List Number")
            doc.add_paragraph(f"Expected result: {case.get('expected_result', '')}")
            doc.add_paragraph("Evidence: " + "; ".join(case.get("evidence", [])))
        for page in pages:
            doc.add_heading(page.section or page.title or "Page", level=1)
            doc.add_paragraph(page.url)
            if page.summary:
                doc.add_heading("Summary", level=2)
                doc.add_paragraph(page.summary)
            doc.add_heading("Structure", level=2)
            doc.add_paragraph(f"Title: {page.title or 'N/A'}")
            doc.add_paragraph(f"Words: {page.word_count} | Headings: {len(page.headings)} | Forms: {len(page.forms)}")
            if page.forms:
                doc.add_heading("Forms", level=2)
                for form in page.forms:
                    doc.add_paragraph(
                        f"Form {form['index']}: {form['method'].upper()} {form['action'] or '(current page)'}",
                        style="List Bullet",
                    )
            if page.generated_test_cases:
                doc.add_heading("Test Cases", level=2)
                for case in page.generated_test_cases:
                    doc.add_paragraph(f"{case['id']} — {case['title']}", style="List Bullet")
        doc.save(path)
        return path

    async def execute_safe_qa_tests(self, test_cases: list, headless=True, max_tests=20) -> list:
        """Execute only non-destructive QA checks; never submits forms or activates arbitrary buttons."""
        results = []
        safe_titles = {
            "Verify form renders correctly", "Verify input controls", "Verify navigation links",
            "Verify page structure", "Review missing heading structure", "Review accessibility findings",
        }
        selected = [case for case in test_cases if case.get("title") in safe_titles][:max_tests]
        if not selected:
            return results
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=headless)
            context = await browser.new_context(viewport={"width": 1460, "height": 1080}, ignore_https_errors=True)
            try:
                for case in selected:
                    started = asyncio.get_running_loop().time()
                    evidence = []
                    screenshot_path = None
                    status = "PASSED"
                    actual = ""
                    try:
                        page = await context.new_page()
                        console_errors = []
                        failed_requests = []
                        page.on("console", lambda msg: console_errors.append(msg.text) if msg.type == "error" and len(console_errors) < 20 else None)
                        page.on("requestfailed", lambda req: failed_requests.append({"url": req.url, "failure": req.failure}) if len(failed_requests) < 20 else None)
                        await page.goto(case["url"], wait_until="domcontentloaded", timeout=30000)
                        await page.wait_for_timeout(300)
                        if case["title"] == "Verify form renders correctly":
                            forms = await page.locator("form").count()
                            status = "PASSED" if forms > 0 else "FAILED"
                            actual = f"Detected {forms} form(s)."
                        elif case["title"] == "Verify input controls":
                            controls = await page.locator("input:not([type=hidden]), select, textarea").count()
                            status = "PASSED" if controls > 0 else "FAILED"
                            actual = f"Detected {controls} visible input control(s)."
                        elif case["title"] == "Verify navigation links":
                            links = await page.locator("a[href]").count()
                            status = "PASSED" if links > 0 else "FAILED"
                            actual = f"Detected {links} navigable link(s)."
                        elif case["title"] in {"Verify page structure", "Review missing heading structure"}:
                            headings = await page.locator("h1, h2, h3, h4, h5, h6").count()
                            status = "PASSED" if headings > 0 else "FAILED"
                            actual = f"Detected {headings} semantic heading(s)."
                        elif case["title"] == "Review accessibility findings":
                            findings = await self.assess_accessibility(page)
                            status = "PASSED" if not findings else "FAILED"
                            actual = f"Detected {len(findings)} accessibility finding(s)."
                            evidence.extend([f"{f.get('severity')}: {f.get('message')}" for f in findings[:10]])
                        evidence.extend([f"Final URL: {page.url}", f"Console errors: {len(console_errors)}", f"Failed requests: {len(failed_requests)}"])
                        screenshot_path = os.path.join(self.debug_dir, f"qa_{case['id']}_{uuid.uuid4().hex[:8]}.png")
                        await page.screenshot(path=screenshot_path, full_page=True)
                        await page.close()
                    except Exception as exc:
                        status = "ERROR"
                        actual = str(exc)
                        evidence.append(f"Execution error: {exc}")
                    results.append({
                        "id": case["id"], "title": case["title"], "url": case["url"],
                        "status": status, "actual_result": actual, "duration_ms": int((asyncio.get_running_loop().time() - started) * 1000),
                        "evidence": evidence, "screenshot_path": screenshot_path or "",
                        "executed_safely": True,
                    })
            finally:
                await context.close()
                await browser.close()
        return results

    async def crawl_and_process(self, max_pages=20, headless=True, progress_cb=None):
        crawl_result = await self.crawl(max_pages, headless=headless, progress_cb=progress_cb)
        regression = self.compare_with_baseline()
        md_path, html_path = self.generate_rag_document()
        docx_path = self.generate_docx_report()
        pages = list(self.content_collection.values())
        return {
            "crawl_summary": {
                k: crawl_result.get(k) for k in [
                    "base_url", "pages_visited_count", "link_count", "content_count",
                    "successful_scrapes", "failed_scrapes", "form_count",
                    "interactive_element_count", "test_case_count", "section_count",
                    "high_risk_pages", "medium_risk_pages", "low_risk_pages",
                    "console_error_count", "failed_request_count", "accessibility_finding_count", "api_request_count", "crawl_duration",
                ]
            },
            "rag_document": {
                "markdown_path": md_path,
                "html_path": html_path,
                "docx_path": docx_path,
                "link_count": len(self.link_collection),
                "content_count": len(self.content_collection),
                "successful_scrapes": len([c for c in pages if c.status.startswith("success")]),
            },
            "pages": [content.model_dump() for content in pages],
            "test_cases": [case for page in pages for case in page.generated_test_cases],
            "regression": regression,
            "execution_results": [],
        }
